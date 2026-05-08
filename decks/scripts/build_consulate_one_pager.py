"""Generate a one-page IndLokal x Consulate of India introduction (EN).

Targets the Consulates General of India in Frankfurt and Munich.

Run:
    python3 decks/scripts/build_consulate_one_pager.py

Output:
    decks/output/IndLokal_Consulate_OnePager.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "IndLokal_Consulate_OnePager.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT = RGBColor(0xFF, 0x8C, 0x00)
SAFFRON = RGBColor(0xFF, 0x99, 0x33)
GREEN = RGBColor(0x13, 0x88, 0x08)
MUTED = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_run(para, text: str, *, bold: bool = False, italic: bool = False,
            size: int = 10, color: RGBColor = DARK):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def tight(para, *, before: int = 0, after: int = 2, line: float = 1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    tight(p, before=4, after=2)
    add_run(p, text.upper(), bold=True, size=10, color=PRIMARY)


def bullet(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    tight(p, after=1, line=1.1)
    if bold_lead:
        add_run(p, bold_lead, bold=True, size=10)
        add_run(p, text, size=10)
    else:
        add_run(p, text, size=10)


def build() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ---- Header band (saffron / white / green nod) --------------------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)
    left.width = Cm(11)
    right.width = Cm(7)
    set_cell_shading(left, "0B3D91")
    set_cell_shading(right, "FF9933")

    p = left.paragraphs[0]
    tight(p, after=0)
    add_run(p, "IndLokal", bold=True, size=20, color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = left.add_paragraph()
    tight(p2, after=0)
    add_run(p2, "Connecting the Indian community in Germany",
            italic=True, size=11, color=RGBColor(0xFF, 0xFF, 0xFF))
    p3 = left.add_paragraph()
    tight(p3, after=0)
    add_run(p3,
            "A digital integration platform for Indians in Baden-Wuerttemberg, Bavaria and Hessen",
            size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp, after=0)
    add_run(rp, "For: Consulate General of India",
            bold=True, size=11, color=RGBColor(0xFF, 0xFF, 0xFF))
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp2, after=0)
    add_run(rp2, "Frankfurt  \u00b7  Muenchen", size=10,
            color=RGBColor(0xFF, 0xFF, 0xFF))
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp3, after=0)
    add_run(rp3, "30-minute introductory meeting",
            italic=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))

    spacer = doc.add_paragraph()
    tight(spacer, after=0)

    # ---- Why we are writing ------------------------------------------
    section_heading(doc, "Why we are writing to the Consulate")
    p = doc.add_paragraph()
    tight(p, after=3)
    add_run(p,
            "The Indian community in Germany has crossed 250,000 (Destatis 2024) "
            "and is the fastest-growing third-country migration to Germany, with "
            "particularly strong concentrations in the consular jurisdictions of "
            "Frankfurt (Hessen, Rheinland-Pfalz, Saarland) and Munich (Bavaria, "
            "Baden-Wuerttemberg). Yet first-time arrivals \u2014 students, IT "
            "professionals, dependants, Ausbildung trainees \u2014 still rely on "
            "fragmented WhatsApp groups, outdated Facebook pages and word of mouth "
            "to find communities, cultural events, temples, language circles and "
            "trusted local resources. The first 90 days in Germany shape long-term "
            "integration outcomes, and that information layer is missing.",
            size=10)

    p = doc.add_paragraph()
    tight(p, after=3)
    add_run(p, "IndLokal closes that gap. ", bold=True, size=10)
    add_run(p,
            "It is a German non-profit-in-formation, built and self-funded by two "
            "India-origin founders living in the Stuttgart region. The platform "
            "is a curated, multilingual (DE/EN, with Hindi/Tamil/Telugu) directory "
            "of verified Indian communities, events and city-specific resources \u2014 "
            "mobile-first, privacy-friendly, and designed to plug into official "
            "structures rather than duplicate them.",
            size=10)

    # ---- What is built today -----------------------------------------
    section_heading(doc, "Status today (honest snapshot)")
    bullet(doc,
           "Web app indlokal.de + native iOS/Android apps fully developed, in private beta. "
           "Public soft launch planned within ~4 weeks (May/June 2026).",
           bold_lead="Platform: ")
    bullet(doc,
           "Verified communities, events and resources schema, multilingual content layer ready.",
           bold_lead="Content: ")
    bullet(doc,
           "Built end-to-end by the two founders, fully self-funded. No external capital, no revenue \u2014 "
           "by design, until the gemeinnuetzig (charitable) entity is in place.",
           bold_lead="Funding so far: ")
    bullet(doc,
           "Gemeinnuetzige UG / gGmbH in formation, registration expected within 4\u20136 weeks.",
           bold_lead="Legal entity: ")
    bullet(doc,
           "Founders: Jaya Prakash Jain (Sindelfingen) and Dhiraj Shah (Waiblingen) \u2014 "
           "Indian citizens, long-term residents in Germany, professional backgrounds in "
           "software engineering and product.",
           bold_lead="Team: ")

    # ---- Why the Consulate matters -----------------------------------
    section_heading(doc, "Why the Consulate's support matters to us")
    bullet(doc,
           "A short letter from the Consulate confirming relevance to the Indian community in "
           "Germany would significantly strengthen our applications to AMIF (EU), ESF+ Baden-Wuerttemberg, "
           "BW Stiftung, Robert Bosch Stiftung and Mercator.",
           bold_lead="Letter of support: ")
    bullet(doc,
           "Cross-listing IndLokal alongside official Consulate notices (community events, festivals, "
           "OCI/passport camps, MEA advisories) so first-time arrivals discover trusted information in one place.",
           bold_lead="Visibility: ")
    bullet(doc,
           "We would be honoured to surface Consulate-organised cultural events, Pravasi Bharatiya Divas, "
           "Republic Day / Independence Day celebrations, and IDY (International Day of Yoga) events to "
           "the right local audience \u2014 free of charge for the Consulate.",
           bold_lead="Two-way channel: ")
    bullet(doc,
           "If discretionary cultural / community-engagement funds (e.g. Indian Council for Cultural Relations, "
           "MEA community welfare allocations, Pravasi Bharatiya support schemes) are accessible, we would "
           "welcome guidance on the right pathway. Funding is welcome but not the primary ask \u2014 visibility and "
           "endorsement come first.",
           bold_lead="Funding (where possible): ")

    # ---- The pilot ---------------------------------------------------
    section_heading(doc, "The first 12 months \u2014 what we will deliver")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Reach", "Communities", "Resources", "Geography"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        pp = c.paragraphs[0]
        add_run(pp, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
    row = table.add_row().cells
    values = [
        "1,500\u20132,500 monthly active Indian users in pilot region (10\u201315% of community); "
        "stretch goal 4,000 by month 12",
        "30+ verified associations, temples, language circles, professional networks "
        "(~40 candidates already mapped)",
        "150+ city-specific resources; 60+ professionally translated (Hindi/Tamil/Telugu)",
        "Stuttgart anchor (months 0\u20133); Munich and Frankfurt beta-live in months 1\u20136; full presence in 6\u20138 German metros by month 12 \u2014 AI-assisted ingestion + curated review makes per-city marginal cost low",
    ]
    for i, v in enumerate(values):
        row[i].text = ""
        pp = row[i].paragraphs[0]
        add_run(pp, v, size=9)

    # ---- The ask ------------------------------------------------------
    section_heading(doc, "What we are asking for")
    p = doc.add_paragraph()
    tight(p, after=2)
    add_run(p, "A 30-minute introductory meeting ", bold=True, size=10)
    add_run(p,
            "with the Consulate \u2014 in person in Frankfurt or Munich, "
            "or via video call \u2014 to:",
            size=10)
    bullet(doc, "demonstrate the platform live (web + mobile, beta build);")
    bullet(doc, "walk through the 12-month pilot scope and KPIs;")
    bullet(doc,
           "discuss a lightweight partnership: a letter of support, optional cross-promotion of "
           "Consulate community events, and \u2014 where Consulate schemes allow \u2014 access to "
           "community-welfare or cultural-engagement funds.")

    p = doc.add_paragraph()
    tight(p, after=2)
    add_run(p,
            "We are not asking the Consulate to commit budget in the first meeting. "
            "We are asking for 30 minutes to introduce ourselves, show the product, and "
            "explore how IndLokal can be useful to the Consulate's own outreach to the Indian "
            "community in Germany.",
            italic=True, size=10, color=MUTED)

    # ---- Footer / contact --------------------------------------------
    spacer = doc.add_paragraph()
    tight(spacer, after=0)

    contact_table = doc.add_table(rows=1, cols=1)
    contact_cell = contact_table.cell(0, 0)
    set_cell_shading(contact_cell, "F2F4F8")
    contact_cell.paragraphs[0].text = ""
    cp = contact_cell.paragraphs[0]
    tight(cp, after=0)
    add_run(cp, "Contact  ", bold=True, size=10, color=PRIMARY)
    add_run(cp, "IndLokal Founding Team  \u00b7  ", size=10)
    add_run(cp, "Jaya Prakash Jain (Sindelfingen)  \u00b7  Dhiraj Shah (Waiblingen)  \u00b7  ", size=10)
    add_run(cp, "hello@indlokal.de", bold=True, size=10, color=PRIMARY)
    add_run(cp, "  \u00b7  indlokal.de", size=10)

    # ---- Footnote ----------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(p, before=4, after=0)
    add_run(p,
            "A short product demo video and screenshots are available on request, "
            "as the public website goes live in May/June 2026.",
            italic=True, size=8, color=MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()

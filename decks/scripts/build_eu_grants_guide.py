"""Generate a layperson-friendly guide to EU grants as a .docx document.

Run:
    python3 decks/scripts/build_eu_grants_guide.py

Output:
    decks/output/EU_Grants_Explained.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "EU_Grants_Explained.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)   # EU-ish blue
ACCENT = RGBColor(0xFF, 0xCC, 0x00)    # EU-ish gold
MUTED = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in [(1, 22), (2, 16), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = PRIMARY
        style.font.bold = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = MUTED

    # decorative rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule.add_run("— ★ —")
    r.font.color.rgb = ACCENT
    r.font.size = Pt(14)
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r = p.add_run(f"💡 {title}\n")
    r.bold = True
    r.font.color.rgb = PRIMARY
    r.font.size = Pt(12)
    p.add_run(body).font.size = Pt(11)
    # light shading via XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF8DC")
    tc_pr.append(shd)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    set_base_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(
        doc,
        "EU Grants, Explained",
        "A plain-English guide for founders, builders, and the merely curious",
    )

    # ---- Intro
    add_h(doc, "1. Why this guide exists", 1)
    add_p(
        doc,
        "Every year the European Union and its member states hand out tens of "
        "billions of euros in grants. That money funds research labs, climate "
        "startups, civic tech, films, farms, and small businesses you have "
        "never heard of. Yet most people — including most founders — have "
        "no idea where to start. The acronyms are intimidating (Horizon, EIC, "
        "EIT, ESF+, ERDF, CEF, LIFE…), the websites read like tax law, and "
        "the application portals feel like a portal to 2003.",
    )
    add_p(
        doc,
        "This guide cuts through that. By the end you should be able to: "
        "(1) explain what an EU grant actually is, (2) tell the major "
        "programmes apart, (3) judge whether your project has a realistic "
        "chance, and (4) know the practical first steps to apply.",
    )

    add_callout(
        doc,
        "Who this is for",
        "Founders thinking about non-dilutive funding, researchers wondering "
        "if their idea fits a call, NGOs hunting for project money, and "
        "anyone who hears 'Horizon Europe' and quietly nods.",
    )

    # ---- What is a grant
    add_h(doc, "2. What is an EU grant, really?", 1)
    add_p(
        doc,
        "A grant is money you do not have to pay back, given to you so that "
        "you carry out a specific project the funder cares about. That last "
        "part matters. Grants are not gifts and they are not investments — "
        "they are co-financed contracts. The EU pays a share (often 50–100%) "
        "of the eligible costs of a project that advances an EU policy goal: "
        "the green transition, digital sovereignty, health, defence, "
        "cohesion between regions, research excellence, and so on.",
    )
    add_h(doc, "Grant vs loan vs equity vs prize", 2)
    add_table(
        doc,
        ["Instrument", "Pay back?", "Give up ownership?", "Typical use"],
        [
            ["Grant", "No", "No", "Risky R&D, public-good projects"],
            ["Loan / blended finance", "Yes", "No", "Scale-up, infrastructure"],
            ["Equity (VC, EIC Fund)", "No", "Yes", "High-growth companies"],
            ["Prize", "No", "No", "Reward for a result already achieved"],
        ],
    )
    add_p(
        doc,
        "EU grants are 'non-dilutive' — they do not take a slice of your "
        "company. That is their superpower. The trade-off is paperwork, "
        "timelines measured in months, and reporting obligations after "
        "you win.",
    )

    # ---- Who runs it
    add_h(doc, "3. Who actually runs EU funding?", 1)
    add_p(
        doc,
        "Think of it as three layers stacked on top of each other:",
    )
    add_bullets(
        doc,
        [
            "EU level — the European Commission and its executive agencies "
            "(CINEA, HaDEA, EISMEA, ERCEA, REA) design and run pan-European "
            "programmes such as Horizon Europe, Digital Europe, the EIC and "
            "Creative Europe.",
            "National level — each member state runs its own ministries and "
            "agencies (e.g. BMBF/Projektträger in Germany, Bpifrance in "
            "France, Enterprise Ireland) that channel both EU and national "
            "money.",
            "Regional / local level — managing authorities for the "
            "structural funds (ERDF, ESF+, Just Transition Fund) sit in "
            "regions and cities. This is where the bulk of cohesion money "
            "actually lands.",
        ],
    )
    add_callout(
        doc,
        "Rule of thumb",
        "If your project is cutting-edge and pan-European → look EU-level. "
        "If it is about jobs, training, or regional development in your "
        "city or Land → look regional. Most founders waste months looking "
        "at the wrong layer.",
    )

    # ---- Major programmes
    add_h(doc, "4. The programmes you should actually know", 1)

    add_h(doc, "Horizon Europe (2021–2027, ~€95.5B)", 2)
    add_p(
        doc,
        "The flagship research and innovation programme. Three pillars:",
    )
    add_bullets(
        doc,
        [
            "Pillar 1 — Excellent Science: ERC grants for individual "
            "researchers, Marie Skłodowska-Curie fellowships, research "
            "infrastructures.",
            "Pillar 2 — Global Challenges & Industrial Competitiveness: "
            "big collaborative projects in health, climate, digital, food, "
            "security. Usually 3–5 partners from different countries.",
            "Pillar 3 — Innovative Europe: home of the European Innovation "
            "Council (EIC) — the part most startups care about.",
        ],
    )

    add_h(doc, "European Innovation Council (EIC)", 2)
    add_p(
        doc,
        "The EU's deep-tech startup engine. Three instruments founders "
        "should remember:",
    )
    add_bullets(
        doc,
        [
            "EIC Pathfinder — early, high-risk research (up to €3–4M, "
            "usually consortia).",
            "EIC Transition — turning a lab result into a prototype "
            "(up to €2.5M).",
            "EIC Accelerator — single-company funding: up to €2.5M grant "
            "plus up to €15M equity from the EIC Fund. The crown jewel and "
            "the hardest to win (success rates around 5%).",
        ],
    )

    add_h(doc, "Digital Europe Programme", 2)
    add_p(
        doc,
        "Less about pure research, more about deployment: AI testing "
        "facilities, European Digital Innovation Hubs, cybersecurity, "
        "high-performance computing, digital skills. Co-financing is "
        "typically 50%, with member states topping up the rest.",
    )

    add_h(doc, "Cohesion funds: ERDF, ESF+, Just Transition Fund", 2)
    add_p(
        doc,
        "Quietly the biggest pot of all (~€392B for 2021–2027). Managed "
        "by regions, not Brussels. ERDF funds infrastructure and SMEs, "
        "ESF+ funds people (training, social inclusion, employment), and "
        "the Just Transition Fund supports regions moving away from "
        "fossil fuels. If you run an SME outside a capital city, this is "
        "often your most realistic source.",
    )

    add_h(doc, "Other programmes worth knowing", 2)
    add_bullets(
        doc,
        [
            "LIFE — environment and climate action.",
            "Creative Europe — culture, media, audiovisual.",
            "Erasmus+ — education, training, youth, sport.",
            "CEF (Connecting Europe Facility) — transport, energy and "
            "digital infrastructure.",
            "Interreg — cross-border cooperation projects.",
            "EU4Health, EDF (defence), CAP (agriculture).",
        ],
    )

    # ---- How to read a call
    add_h(doc, "5. How to read a call for proposals without crying", 1)
    add_p(
        doc,
        "Every grant opportunity is published as a 'call' on the EU "
        "Funding & Tenders Portal. Calls look intimidating but follow a "
        "predictable pattern. Learn the anatomy and you can scan one in "
        "ten minutes.",
    )
    add_numbered(
        doc,
        [
            "Topic ID and title — the unique handle (e.g. HORIZON-EIC-"
            "2025-ACCELERATOROPEN-01). Bookmark it.",
            "Type of action — RIA (research), IA (innovation), CSA "
            "(coordination), Lump Sum, etc. Determines funding rate.",
            "Expected outcomes & scope — what the Commission wants the "
            "world to look like after the project. Your proposal must "
            "speak to these in their language.",
            "Eligibility — who can apply, from where, alone or in "
            "consortium, minimum number of partners.",
            "Budget and funding rate — total call budget, indicative "
            "project size, share the EU pays.",
            "Deadlines — single-stage, two-stage, or rolling.",
            "Evaluation criteria — almost always Excellence, Impact, and "
            "Implementation, each scored out of 5.",
        ],
    )
    add_callout(
        doc,
        "Insider tip",
        "The 'Expected Outcomes' section is the brief. If your project "
        "does not naturally produce those outcomes, do not contort it — "
        "find a different call. Forcing a fit is the #1 reason proposals "
        "fail.",
    )

    # ---- Eligibility
    add_h(doc, "6. Are you even eligible?", 1)
    add_p(
        doc,
        "Most EU programmes require you to be legally established in an "
        "EU member state or an 'associated country' (Norway, Iceland, "
        "Israel, the UK for parts of Horizon, Ukraine, and others — the "
        "list shifts). Specific rules vary, but the common gates are:",
    )
    add_bullets(
        doc,
        [
            "You have a legal entity (company, NGO, university, "
            "municipality). Sole traders sometimes qualify, often not.",
            "You are registered in the EU Funding & Tenders Portal and "
            "have a PIC (Participant Identification Code).",
            "For SME-specific calls, you meet the EU SME definition: "
            "<250 employees, ≤€50M turnover or ≤€43M balance sheet, and "
            "not majority-owned by a larger company.",
            "You are not in financial distress, not under EU sanctions, "
            "and have no unresolved fraud findings.",
        ],
    )

    # ---- Application process
    add_h(doc, "7. The application process, step by step", 1)
    add_numbered(
        doc,
        [
            "Find the right call. Use the Funding & Tenders Portal "
            "search, filter by programme and deadline. Read at least "
            "three winning proposals from previous rounds if you can find "
            "them.",
            "Register your organisation. Get a PIC, upload statutes, "
            "financial statements, VAT, etc. Do this weeks before the "
            "deadline — validation can take time.",
            "Build the consortium (if needed). Many calls require 3+ "
            "partners from 3+ countries. Use partner-search tools, "
            "national contact points (NCPs), and existing networks.",
            "Draft Part A (administrative) and Part B (the technical "
            "proposal). Part B is where you win or lose. Stick to the "
            "page limit; evaluators stop reading at the limit.",
            "Internal review. Have someone outside your bubble read it. "
            "If they cannot summarise your impact in one sentence, "
            "neither will the evaluator.",
            "Submit before the deadline. The portal closes at 17:00 "
            "Brussels time, sharp. Submit at least a day early — the "
            "portal famously melts in the final hour.",
            "Evaluation. Independent expert evaluators score you on "
            "Excellence, Impact, Implementation. You get an Evaluation "
            "Summary Report (ESR) — read it carefully even if you win.",
            "Grant Agreement Preparation (GAP). If selected, you "
            "negotiate the contract, budget, and start date. This takes "
            "2–6 months.",
            "Implementation and reporting. Periodic reports, financial "
            "statements, audits. Keep clean timesheets from day one.",
        ],
    )

    # ---- Costs and money
    add_h(doc, "8. How the money actually flows", 1)
    add_p(
        doc,
        "EU grants reimburse 'eligible costs' at a defined funding rate. "
        "Typical eligible cost categories:",
    )
    add_bullets(
        doc,
        [
            "Personnel — salaries based on actual hours worked on the "
            "project. Hence the timesheet obsession.",
            "Subcontracting — limited; the EU prefers you do the work "
            "yourself.",
            "Travel & subsistence — capped per diems.",
            "Equipment — usually only the depreciation during the "
            "project, not the full purchase price.",
            "Other goods, works and services — consumables, software, "
            "dissemination, audits.",
            "Indirect costs — a flat 25% on top of direct costs in most "
            "Horizon actions.",
        ],
    )
    add_p(
        doc,
        "Funding rates: 100% for research actions and most non-profits; "
        "70% for innovation actions when you are a for-profit company "
        "(100% if non-profit). Lump-sum grants are increasingly common — "
        "you get a fixed amount per work package on completion, no "
        "timesheets.",
    )
    add_callout(
        doc,
        "Cash-flow warning",
        "You typically receive a 'pre-financing' payment (often ~40%) "
        "after signing, then interim payments after each reporting "
        "period, then a final payment after the project ends. Plan for "
        "12–18 months of working capital. Many startups die of "
        "cash-flow, not rejection.",
    )

    # ---- Common mistakes
    add_h(doc, "9. The mistakes everyone makes (so you don't have to)", 1)
    add_bullets(
        doc,
        [
            "Writing the proposal you want to write, not the one the "
            "call asks for.",
            "Treating impact as an afterthought. Impact is half the "
            "score; start with it.",
            "Vague work plan. Evaluators want concrete deliverables, "
            "milestones, a Gantt chart, and a credible risk register.",
            "Ignoring dissemination, communication, and exploitation. "
            "These are mandatory sections, not filler.",
            "Underestimating ethics, gender, open science, and 'do no "
            "significant harm' requirements.",
            "Submitting in the last hour. Portal outages have killed "
            "more proposals than weak ideas.",
            "No commercialisation story (for innovation actions). The "
            "EU wants to see real market uptake, not just a paper.",
        ],
    )

    # ---- Help
    add_h(doc, "10. Where to get help — for free", 1)
    add_bullets(
        doc,
        [
            "National Contact Points (NCPs) — every country has one per "
            "programme. They review draft proposals, host info days, and "
            "the service is free.",
            "Enterprise Europe Network (EEN) — local advisors for SMEs.",
            "European Digital Innovation Hubs (EDIHs) — regional "
            "support for digital transformation projects.",
            "EIC Community platform and the Funding & Tenders Portal "
            "partner-search.",
            "Your regional managing authority — for ERDF / ESF+ money "
            "this is the only door.",
            "Universities and research organisations — they have grant "
            "offices that have done this hundreds of times.",
        ],
    )
    add_p(
        doc,
        "Paid consultants exist and can be excellent, but fees range "
        "from a few thousand euros to 10–15% success fees. Talk to NCPs "
        "first — you may not need one.",
    )

    # ---- Realism
    add_h(doc, "11. A reality check before you start", 1)
    add_bullets(
        doc,
        [
            "Success rates are low. EIC Accelerator: ~5%. Horizon "
            "collaborative calls: 10–20%. ERC: ~13%. A rejection is the "
            "default outcome; treat each attempt as a draft.",
            "Time investment is real. A serious Horizon proposal is "
            "200–400 person-hours. An EIC Accelerator full proposal is "
            "more.",
            "Timelines are slow. From call publication to money in the "
            "bank: typically 9–14 months.",
            "Reporting is mandatory. Plan ~5–10% of project effort for "
            "admin and reporting.",
            "Audits happen. Years later. Keep every receipt, contract "
            "and timesheet for at least 5 years after the final payment.",
        ],
    )
    add_callout(
        doc,
        "When grants are a bad fit",
        "If you need money in 3 months, if your project pivots every "
        "quarter, or if the topic does not match any policy priority — "
        "skip grants. Use revenue, angels, or a bank loan. Grants are "
        "patient, slow, public money for projects with a clear public "
        "benefit.",
    )

    # ---- Glossary
    add_h(doc, "12. Glossary of acronyms you will meet", 1)
    add_table(
        doc,
        ["Acronym", "Meaning"],
        [
            ["EC", "European Commission"],
            ["EIC", "European Innovation Council"],
            ["EIT", "European Institute of Innovation & Technology"],
            ["ERC", "European Research Council"],
            ["ERDF", "European Regional Development Fund"],
            ["ESF+", "European Social Fund Plus"],
            ["CEF", "Connecting Europe Facility"],
            ["LIFE", "EU programme for Environment & Climate Action"],
            ["MSCA", "Marie Skłodowska-Curie Actions"],
            ["NCP", "National Contact Point"],
            ["PIC", "Participant Identification Code"],
            ["RIA / IA / CSA", "Research / Innovation / Coordination & Support Action"],
            ["TRL", "Technology Readiness Level (1 = idea, 9 = on the market)"],
            ["WP", "Work Package (and also Work Programme — context!)"],
            ["GA", "Grant Agreement"],
            ["ESR", "Evaluation Summary Report"],
        ],
    )

    # ---- First steps
    add_h(doc, "13. Your first 7 days", 1)
    add_numbered(
        doc,
        [
            "Open the EU Funding & Tenders Portal and create an EU Login.",
            "Search 2–3 calls that match your topic. Read them fully.",
            "Identify your National Contact Point and email them with a "
            "one-page summary of your idea.",
            "Find one previously funded project on CORDIS that resembles "
            "yours. Read the public summary.",
            "Sketch your project on one page: problem, solution, who "
            "benefits, expected impact, partners, budget order of "
            "magnitude.",
            "Decide honestly: grant, loan, equity, or revenue? Do not "
            "default to grants because they sound free.",
            "If you proceed, register your organisation for a PIC today. "
            "It is the longest-lead item.",
        ],
    )

    # ---- Closing
    add_h(doc, "14. Final word", 1)
    add_p(
        doc,
        "EU grants are not magic money and they are not a scam. They are "
        "a serious instrument for serious projects with public benefit. "
        "Approach them like you would a Series A: with a sharp story, "
        "the right partners, a credible plan, and the patience to "
        "iterate. Win or lose, you will come out with a better-defined "
        "project — and that, more than the cheque, is often the real "
        "prize.",
    )
    add_p(doc, "Good luck. Europe is rooting for you.")

    # ---- Disclaimer
    doc.add_paragraph()
    disc = doc.add_paragraph()
    r = disc.add_run(
        "Disclaimer: This guide is for general orientation. Programme "
        "names, budgets, rules and deadlines change frequently. Always "
        "check the official EU Funding & Tenders Portal "
        "(ec.europa.eu/info/funding-tenders) and your National Contact "
        "Point for current information before applying."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()

"""Generate the IndLokal grant strategy memo as a .docx document.

Run:
    python3 decks/scripts/build_indlokal_grant_strategy.py

Output:
    decks/output/IndLokal_Grant_Strategy.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "IndLokal_Grant_Strategy.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT = RGBColor(0xFF, 0x8C, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
GOOD = RGBColor(0x1B, 0x7A, 0x3E)
BAD = RGBColor(0xB0, 0x2A, 0x2A)


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [(1, 20), (2, 15), (3, 12)]:
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = PRIMARY
        s.font.bold = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = MUTED

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = rule.add_run("— ★ —")
    r3.font.color.rgb = ACCENT
    r3.font.size = Pt(13)
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_callout(doc: Document, title: str, body: str, fill: str = "FFF8DC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}\n")
    r.bold = True
    r.font.color.rgb = PRIMARY
    r.font.size = Pt(12)
    p.add_run(body).font.size = Pt(11)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def add_checklist(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph()
        r = p.add_run("☐  ")
        r.font.size = Pt(12)
        r.bold = True
        p.add_run(it).font.size = Pt(11)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    set_base_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(
        doc,
        "IndLokal × EU/DE Grant Landscape",
        "Strategic fit memo — where we belong, what to fix, where to apply, how to apply",
    )

    # --- Section 1
    add_h(doc, "1. The honest one-line fit", 1)
    add_p(
        doc,
        "IndLokal is digital integration infrastructure for third-country "
        "nationals (TCNs) in Germany. That is exactly the language of one "
        "EU programme (AMIF) and a stack of German federal, Land and "
        "foundation pots. It is not a research, deep-tech or pan-European "
        "consortium project — so most of the headline EU programmes "
        "(Horizon, EIC, Digital Europe) are a poor fit. Don't chase them.",
    )
    add_callout(
        doc,
        "Bottom line",
        "Optimise for AMIF, ESF+ Baden-Württemberg, Land BW integration "
        "funds, Stadt Stuttgart, and the Stuttgart-region foundations "
        "(Mercator, Bosch, BW Stiftung). Treat everything else as "
        "opportunistic.",
    )
    add_callout(
        doc,        "Reality check before reading on",
        "Status today: platform built, in private beta, soft public launch "
        "within ~4 weeks. No legal entity yet (still deciding between "
        "gUG, e.V., gGmbH). No external capital, no revenue. This "
        "limits what we can apply for in month 1\u20132 \u2014 most public "
        "funders and foundations require a legal applicant. The plan "
        "below sequences entity setup and soft launch BEFORE the bigger "
        "asks.",
        fill="FFE8E0",
    )
    add_callout(
        doc,        "How to think about the €150k headline number",
        "No Stuttgart-region foundation will hand a first-time grantee "
        "€150k in one cheque. The realistic path is either (a) one large "
        "AMIF / ESF+ award, or (b) stacking ~3 smaller awards: e.g. "
        "€15k Stadt Stuttgart + €50k Bosch + €85k BW Stiftung = €150k. "
        "The deck shows three explicit tiers — €85k lean (single "
        "funder), €150k full (single large or stacked), €250k Region "
        "Stuttgart (Stuttgart + Sindelfingen + Waiblingen).",
    )

    # --- Section 2
    add_h(doc, "2. Programme-by-programme fit", 1)

    add_h(doc, "Tier A — apply here first (high fit, realistic to win)", 2)
    add_table(
        doc,
        ["Programme", "Owner", "Fit", "Typical size", "Why"],
        [
            ["AMIF — National Programme DE", "BAMF",
             "★★★★★", "€100k–€500k",
             "TCN integration is the Specific Objective. Multi-language + civic participation + public-sector partnership = textbook AMIF."],
            ["ESF+ Baden-Württemberg", "Min. f. Soziales BW",
             "★★★★★", "€50k–€300k",
             "Social inclusion, integration, digital access. Land-level decisions, closer to home."],
            ["Land BW Integrationsförderung / 'Gemeinsam in Vielfalt III'",
             "Min. f. Soziales BW",
             "★★★★★", "€20k–€150k",
             "Direct integration funding for projects in BW cities."],
            ["Stadt Stuttgart Integrationsfonds", "Stadt Stuttgart",
             "★★★★★", "€5k–€50k",
             "Small but fast. Builds the local credibility you need for bigger asks. Often the first 'yes'."],
            ["Stiftung Mercator — Teilhabe & Zusammenhalt", "Foundation",
             "★★★★☆", "€100k–€500k+",
             "Already in the deck. Strong fit, multi-year possible."],
            ["Robert Bosch Stiftung — Migration & Inklusion", "Foundation",
             "★★★★☆", "€50k–€300k",
             "Stuttgart-headquartered. Local angle helps."],
            ["BW Stiftung — Integration / Gesellschaft & Kultur", "Foundation",
             "★★★★☆", "€30k–€200k",
             "Land-level foundation, prefers BW projects."],
        ],
    )

    add_h(doc, "Tier B — apply selectively", 2)
    add_table(
        doc,
        ["Programme", "Fit", "Notes"],
        [
            ["CERV (Citizens, Equality, Rights and Values)", "★★★☆☆",
             "EU strand on Union Values / civic engagement. Usually needs consortium. Frame: civic participation of new EU residents."],
            ["Prototype Fund (BMBF + OKF)", "★★★☆☆",
             "€95k for open-source civic tech. Requires releasing a meaningful module as OSS."],
            ["Hertie Stiftung — Demokratie stärken", "★★★☆☆",
             "Civic participation angle."],
            ["Schöpflin Stiftung", "★★★☆☆",
             "Civil society digital infrastructure."],
            ["Allianz Foundation", "★★★☆☆",
             "EU-wide, solidarity / pluralism."],
            ["Körber-Stiftung", "★★★☆☆",
             "Migration & city-society projects (Hamburg-leaning, but open)."],
            ["Google.org Impact Challenge / Meta Community Accelerator",
             "★★★☆☆",
             "Need nonprofit status (gGmbH or e.V.) for most cycles."],
        ],
    )

    add_h(doc, "Tier C — wrong fit, do not apply", 2)
    add_bullets(
        doc,
        [
            "EIC Accelerator / Pathfinder / Transition — deep-tech only. Will be desk-rejected.",
            "Horizon Europe Pillar 2 — needs 3+ partners across 3+ countries; research focus; novelty in method, not application.",
            "Digital Europe — EU digital sovereignty (HPC, AI, cyber, EDIHs). Wrong space.",
            "CEF Digital — telecom / cross-border infrastructure. Wrong scale.",
        ],
    )

    # --- Section 3
    add_h(doc, "3. Things to check & fix BEFORE applying", 1)

    add_h(doc, "3.1 Legal form is your biggest single lever", 2)
    add_p(
        doc,
        "A standard GmbH is fine for some funders (Bosch, Mercator, AMIF "
        "can fund GmbHs), but locks you out of others:",
    )
    add_bullets(
        doc,
        [
            "Google Ad Grants (~€10k/month free Google Ads) — needs registered nonprofit.",
            "Google Workspace / Microsoft / Apple / GitHub nonprofit programmes.",
            "Many Land programmes that require gemeinnützig status.",
            "Reduced co-financing rates in several public calls.",
        ],
    )
    add_callout(
        doc,
        "Recommended action",
        "Set up a parallel gGmbH or e.V. that receives grant funding and "
        "contracts technology services from the GmbH at arm's length. "
        "Cost ~€2–5k, a few weeks via a Steuerberater specialised in "
        "Gemeinnützigkeit. Many social ventures use this exact structure.",
    )

    add_h(doc, "3.2 Co-financing & cash flow", 2)
    add_table(
        doc,
        ["Source", "EU/funder share", "Co-financing"],
        [
            ["AMIF", "up to 75%",
             "25%, often covered by Land or municipal partner"],
            ["ESF+ BW", "40–60%",
             "Often covered by Land or municipality"],
            ["Foundations (Mercator, Bosch, BW Stiftung)", "up to 100%",
             "Usually none required"],
            ["Stadt Stuttgart Integrationsfonds", "up to 80–100%",
             "Sometimes co-funding from partner association"],
        ],
    )
    add_callout(
        doc,
        "Cash-flow reality",
        "Pre-financing is typically 20–40%. You need ~6 months of "
        "working capital between signing and the first interim payment. "
        "Plan for it — many small grantees die of cash-flow, not "
        "rejection.",
    )

    add_h(doc, "3.3 Mandatory paperwork checklist", 2)
    add_checklist(
        doc,
        [
            "PIC code on the EU Funding & Tenders Portal (free, slow — start now)",
            "EU Login for both founders",
            "DSGVO concept: DPIA for community/event data, AVV with hosting providers, ROPA (Verarbeitungsverzeichnis)",
            "Accessibility statement per BFSG (mandatory since June 2025 for digital services in DE)",
            "Impressum + Datenschutzerklärung legally clean and reviewed",
            "Letters of intent / MoUs from Stadt Stuttgart Welcome Center, IHK Stuttgart, ≥1 community partner, ≥1 university",
            "Independent evaluator with a signed offer letter (target: freelance evaluation consultant or a chair at Uni Stuttgart / Hochschule Esslingen; ifo / BAMF-FZ as research-link aspirations only)",
            "Audited or reviewed financials for the last 1–2 years",
            "Project bank account separated from operating account (required for AMIF reporting)",
        ],
    )

    add_h(doc, "3.4 Re-skin the deck per funder", 2)
    add_table(
        doc,
        ["Funder", "Re-frame the deck around…"],
        [
            ["AMIF / BAMF",
             "TCN integration in the first 90 days, civic participation, multi-language access. Drop startup language."],
            ["Stadt Stuttgart",
             "Hyper-local: districts, named partners, named events, named community leaders. Very concrete."],
            ["Mercator / Bosch",
             "Social cohesion, evidence-based, scalable model, independent evaluation. Avoid name-dropping BAMF-FZ / ifo unless an email confirmation exists \u2014 use freelance evaluator or university chair as the primary commitment."],
            ["BW Stiftung",
             "Baden-Württemberg first; explicit alignment with the Land integration plan."],
            ["Prototype Fund",
             "Open-source civic infrastructure for migrant communities. Emphasise the OSS component."],
        ],
    )

    # --- Section 4
    add_h(doc, "4. Application playbook — what to do, in order", 1)

    add_h(doc, "Next 2 weeks", 2)
    add_numbered(
        doc,
        [
            "Register PIC on the EU Funding & Tenders Portal — slowest item, start today.",
            "Email Stadt Stuttgart Abteilung Integration: request a 30-min meeting and share the 1-pager. Their Letter of Support is the single most valuable document you can collect.",
            "Email BAMF Regionalkoordinator BW for AMIF national programme guidance and the current call schedule.",
            "Email the BW Ministerium für Soziales referat for Integrationsförderung — ask which open call best fits a 12-month Stuttgart pilot.",
            "Talk to a Steuerberater about gGmbH/e.V. setup. Get a written cost + timeline.",
        ],
    )

    add_h(doc, "Weeks 3–6", 2)
    add_numbered(
        doc,
        [
            "Pick your first two applications. Recommended: (a) Stadt Stuttgart Integrationsfonds — small, fast, builds CV; (b) Stiftung Mercator OR Robert Bosch Stiftung — bigger, longer cycle.",
            "Sign MoUs with 3 community organisations + 1 university international office + Stadt Stuttgart Welcome Center.",
            "Get the independent evaluator's signed offer letter.",
            "Draft application v1, then submit internally to a friendly NCP reviewer.",
        ],
    )

    add_h(doc, "Weeks 7–12", 2)
    add_numbered(
        doc,
        [
            "Submit Stadt Stuttgart application (rolling or quarterly deadlines).",
            "Submit foundation application (Mercator/Bosch — usually rolling, decisions in 3–6 months).",
            "Prepare AMIF / ESF+ application for the next published call (BAMF + L-Bank BW pages).",
        ],
    )

    add_h(doc, "Months 4–6", 2)
    add_numbered(
        doc,
        [
            "Apply for Google for Nonprofits (after gGmbH/e.V. is registered) → unlocks Ad Grants + free Workspace.",
            "Apply for Microsoft for Nonprofits → Azure credits (could replace your hosting line item).",
            "Apply to Prototype Fund if you can carve out an OSS module.",
        ],
    )

    # --- Section 5
    add_h(doc, "5. Advertising, marketing & in-kind channels", 1)
    add_p(
        doc,
        "Not grants, but they materially reduce burn and add reach. "
        "Most require nonprofit status — another argument for the "
        "gGmbH/e.V. structure.",
    )
    add_table(
        doc,
        ["Channel", "What you get", "Requirement"],
        [
            ["Google Ad Grants", "~€9k/month free Google Ads", "Registered nonprofit, website meets policy"],
            ["Google Workspace for Nonprofits", "Free Workspace Business Standard", "Same"],
            ["Microsoft for Nonprofits", "Azure credits, M365 free/discounted", "Same"],
            ["Meta Nonprofit programs", "Reduced fundraising fees, training", "Same"],
            ["Apple App Store nonprofit fee waiver", "App fees waived", "Same"],
            ["GitHub for Nonprofits", "Free Team plan", "Same"],
            ["Notion / Slack / Figma for Nonprofits", "Free or discounted plans", "Same"],
            ["Stadt Stuttgart Plakatflächen", "Free poster space for civic events", "Partnership with Stadt"],
            ["VVS / SSB co-promotion", "Posters in trams / U-Bahn", "Civic partnership"],
            ["Indo-German Chamber (IGCC) newsletter", "Reach into Indian-corporate diaspora", "Membership / content partnership"],
            ["Indian Consulate Munich newsletter & events", "Direct reach into target audience", "Already a named partner"],
            ["University international offices", "Direct reach into Indian student cohort", "MoU"],
            ["DAAD alumni network", "Reach into highly-skilled returnees", "Light partnership"],
        ],
    )
    add_callout(
        doc,
        "Underused bucket: corporate sponsorship",
        "Indian IT majors (TCS, Infosys, Wipro, HCL, Tech Mahindra) all "
        "have Stuttgart-region offices and CSR budgets earmarked for "
        "diaspora / integration. So do German firms with large Indian "
        "workforces (Bosch, Mercedes-Benz, Porsche, SAP, Allianz). A "
        "€5k–€25k sponsorship from any one is faster than any grant — "
        "and counts as co-financing for AMIF / ESF+.",
    )

    # --- Section 6
    add_h(doc, "6. What to NOT do", 1)
    add_bullets(
        doc,
        [
            "Don't apply to EIC Accelerator — 200 hours wasted, desk-rejected.",
            "Don't try to lead a Horizon consortium. Maybe join one as a community partner.",
            "Don't promise the same KPIs to every funder — they talk to each other (especially BW foundations).",
            "Don't position as 'startup' to public funders or as 'NGO' to commercial sponsors. Use the integration-infrastructure framing for both.",
            "Don't accept any grant that demands open-ended IP transfer or perpetual free-service obligations without legal review.",
        ],
    )

    # --- Section 7
    add_h(doc, "7. Where to send the existing deck — scorecard", 1)
    add_table(
        doc,
        ["Apply now (deck is ready)", "Apply after small tailoring", "Don't bother"],
        [
            ["Stadt Stuttgart Integrationsfonds", "AMIF national programme (BAMF)", "EIC Accelerator"],
            ["Robert Bosch Stiftung", "ESF+ BW", "Horizon Europe"],
            ["Stiftung Mercator", "BW Stiftung", "Digital Europe"],
            ["BW Stiftung discretionary", "CERV (with consortium)", "CEF Digital"],
            ["Prototype Fund (with OSS angle)", "—", "Creative Europe"],
        ],
    )

    # --- Section 8
    add_h(doc, "8. Single most important next action", 1)
    add_callout(
        doc,
        "Do this week",
        "Book a 30-minute meeting with the Stadt Stuttgart Abteilung "
        "Integration. Their letter of support is required (or strongly "
        "recommended) by AMIF, ESF+ BW, BW Stiftung, Bosch and Mercator. "
        "One meeting unlocks five applications.",
        fill="E8F0FF",
    )

    # --- Closing
    doc.add_paragraph()
    disc = doc.add_paragraph()
    r = disc.add_run(
        "Working memo prepared for the IndLokal founding team. "
        "Programme names, eligibility and call schedules change "
        "frequently — verify on the EU Funding & Tenders Portal, "
        "BAMF, L-Bank BW and each foundation's site before applying."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()

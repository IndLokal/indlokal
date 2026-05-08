"""Generate the IndLokal Consulate Outreach Pack:
- Cold email (English) to Consulate General of India, Frankfurt
- Cold email (English) to Consulate General of India, Munich
- Follow-up email (after 7 days, no reply)
- 30-minute meeting agenda
- Product demo video script (90-120 seconds)
- Screenshot pack checklist (for attachments)
- Pre-meeting briefing notes (talking points, anticipated questions)

Run:
    python3 decks/scripts/build_consulate_outreach_pack.py

Output:
    decks/output/IndLokal_Consulate_Outreach_Pack.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "IndLokal_Consulate_Outreach_Pack.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT = RGBColor(0xFF, 0x8C, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_run(para, text: str, *, bold: bool = False, italic: bool = False,
            size: int = 11, color: RGBColor = DARK, mono: bool = False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Consolas" if mono else "Calibri"
    return r


def tight(para, *, before: int = 0, after: int = 4, line: float = 1.2):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_title(doc: Document, text: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    tight(p, after=2)
    add_run(p, text, bold=True, size=22, color=PRIMARY)
    if subtitle:
        p2 = doc.add_paragraph()
        tight(p2, after=10)
        add_run(p2, subtitle, italic=True, size=12, color=MUTED)


def add_h(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    if level == 1:
        tight(p, before=14, after=4)
        add_run(p, text, bold=True, size=15, color=PRIMARY)
    else:
        tight(p, before=8, after=2)
        add_run(p, text, bold=True, size=12, color=ACCENT)


def add_p(doc: Document, text: str, *, italic: bool = False):
    p = doc.add_paragraph()
    tight(p, after=4)
    add_run(p, text, italic=italic, size=11)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    tight(p, after=2, line=1.15)
    add_run(p, text, size=11)


def add_callout(doc: Document, title: str, body: str, fill: str = "EAF1FB"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.paragraphs[0].text = ""
    p1 = cell.paragraphs[0]
    tight(p1, after=2)
    add_run(p1, title, bold=True, size=11, color=PRIMARY)
    p2 = cell.add_paragraph()
    tight(p2, after=2)
    add_run(p2, body, size=10)
    spacer = doc.add_paragraph()
    tight(spacer, after=2)


def add_email_block(doc: Document, lines: list[str]):
    """Render an email body in a shaded mono block for easy copy-paste."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    cell.paragraphs[0].text = ""
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        tight(p, after=2, line=1.25)
        add_run(p, line, size=10.5)
    spacer = doc.add_paragraph()
    tight(spacer, after=4)


def build() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_title(
        doc,
        "IndLokal \u2014 Consulate Outreach Pack",
        "Cold emails, follow-ups, meeting agenda, demo script, screenshot checklist and briefing notes "
        "for outreach to the Consulates General of India in Frankfurt and Munich",
    )

    add_callout(
        doc,
        "How to use this pack",
        "1. Send the relevant city email below (Frankfurt OR Munich) with three attachments: "
        "(a) IndLokal_Consulate_OnePager.docx, (b) the 90-second demo video MP4, "
        "(c) IndLokal_Screenshots.pdf. 2. If no reply in 7 working days, send the follow-up. "
        "3. When the meeting is confirmed, use the agenda + briefing notes to prepare. "
        "4. Goal of the meeting: a verbal yes for a letter of support. Funding talk is secondary.",
        fill="FFE8E0",
    )

    # ------------------------------------------------------------------
    # Section 1 \u2014 Frankfurt email
    # ------------------------------------------------------------------
    add_h(doc, "1. Cold email \u2014 Consulate General of India, Frankfurt")

    add_p(doc, "To: cons.frankfurt@mea.gov.in (general); cc: hoc.frankfurt@mea.gov.in (Head of Chancery)", italic=True)
    add_p(doc, "Subject: IndLokal \u2014 a digital integration platform for the Indian community in Germany; request for a 30-min introduction", italic=True)

    add_email_block(doc, [
        "Respected Sir / Madam,",
        "",
        "I write on behalf of IndLokal, an India-origin founder team based in the "
        "Stuttgart region, building a digital integration platform for the Indian "
        "community in Germany. The platform \u2014 a curated, multilingual directory "
        "of verified Indian communities, cultural events, temples, language circles "
        "and city-specific resources \u2014 is fully developed and goes into public "
        "soft launch in May/June 2026.",
        "",
        "We are writing to the Consulate General in Frankfurt because Hessen, "
        "Rheinland-Pfalz and Saarland together host one of the fastest-growing "
        "Indian communities in Germany, and because much of what we are building "
        "complements the Consulate's own community outreach \u2014 cultural events, "
        "OCI/passport camps, festival celebrations, and information for first-time "
        "arrivals.",
        "",
        "We would be honoured if the Consulate could spare 30 minutes \u2014 in person "
        "in Frankfurt or via video call \u2014 in May or early June 2026, to:",
        "",
        "  \u2022  introduce ourselves (Indian citizens, long-term residents in Germany),",
        "  \u2022  demonstrate the platform live (web + iOS/Android beta),",
        "  \u2022  walk through the 12-month pilot plan and KPIs,",
        "  \u2022  and discuss how a short letter of support from the Consulate, plus "
        "optional cross-promotion of Consulate-organised community events through "
        "the platform, could meaningfully strengthen the initiative.",
        "",
        "Where the Consulate's discretionary community-welfare or cultural-engagement "
        "channels (ICCR, MEA community welfare, Pravasi Bharatiya schemes) permit, "
        "we would also welcome guidance on the appropriate pathway. However, our "
        "primary request in this first meeting is simply visibility and endorsement, "
        "not financial commitment.",
        "",
        "For your reference, please find attached:",
        "",
        "  1. IndLokal \u2014 Consulate One-Pager (PDF/DOCX, 1 page)",
        "  2. A 90-second product demo video (the website is not yet public)",
        "  3. A short pack of screenshots from the iOS / Android / web app",
        "",
        "We can travel to Frankfurt at the Consulate's convenience and would be "
        "happy to align with any community event or open-house schedule.",
        "",
        "With respect and Jai Hind,",
        "",
        "Jaya Prakash Jain  \u00b7  Co-founder, IndLokal",
        "Schoeneberger Weg 14, 71065 Sindelfingen",
        "M: +49 [number]   \u00b7   E: jaya@indlokal.de",
        "",
        "Dhiraj Shah  \u00b7  Co-founder, IndLokal",
        "Bluetenaecker 17/1, 71332 Waiblingen",
        "M: +49 [number]   \u00b7   E: dhiraj@indlokal.de",
        "",
        "Web (live from May/June 2026): indlokal.de",
        "Legal entity: gemeinnuetzige UG (haftungsbeschraenkt) i.G.",
    ])

    # ------------------------------------------------------------------
    # Section 2 \u2014 Munich email
    # ------------------------------------------------------------------
    add_h(doc, "2. Cold email \u2014 Consulate General of India, Munich")

    add_p(doc, "To: cons.munich@mea.gov.in (general); cc: hoc.munich@mea.gov.in (Head of Chancery)", italic=True)
    add_p(doc, "Subject: IndLokal \u2014 a digital integration platform for Indians in Bavaria and Baden-Wuerttemberg; request for a 30-min introduction", italic=True)

    add_email_block(doc, [
        "Respected Sir / Madam,",
        "",
        "I write on behalf of IndLokal, an India-origin founder team based in the "
        "Stuttgart region (Sindelfingen and Waiblingen), building a digital "
        "integration platform for the Indian community in Germany. The platform \u2014 "
        "a curated, multilingual directory of verified Indian communities, cultural "
        "events, temples, language circles and city-specific resources \u2014 is fully "
        "developed and enters public soft launch in May/June 2026.",
        "",
        "We approach the Consulate General in Munich because Bavaria and "
        "Baden-Wuerttemberg together host the largest Indian community in Germany "
        "and fall within the Consulate's jurisdiction. Our Stuttgart anchor "
        "launches in May/June 2026, and Munich is in our very first wave of "
        "city launches \u2014 beta-live within 1\u20133 months of soft launch, not in \"year 2\". "
        "AI-assisted content ingestion and a city-agnostic platform mean per-city "
        "marginal cost is low; the real bottleneck is trust and partnerships, which "
        "is precisely why we are seeking the Consulate's endorsement now \u2014 before "
        "we land in your jurisdiction, not after.",
        "",
        "We would be deeply grateful for 30 minutes of the Consulate's time \u2014 in "
        "person in Munich or via video call \u2014 in May or early June 2026, to:",
        "",
        "  \u2022  introduce ourselves and the founding team,",
        "  \u2022  demonstrate the platform live (web + iOS/Android beta),",
        "  \u2022  share the 12-month pilot plan, KPIs and partnership approach,",
        "  \u2022  and discuss how a letter of support from the Consulate, plus optional "
        "cross-listing of Consulate-organised cultural and community events on the "
        "platform, could significantly strengthen the initiative.",
        "",
        "Where the Consulate's discretionary community-welfare or cultural channels "
        "(ICCR, MEA community welfare, Pravasi Bharatiya schemes) permit, we would "
        "also welcome guidance. Our primary ask in this first meeting is, however, "
        "visibility and an endorsement \u2014 not financial commitment.",
        "",
        "Attached for your reference:",
        "",
        "  1. IndLokal \u2014 Consulate One-Pager (PDF/DOCX, 1 page)",
        "  2. A 90-second product demo video (the website is not yet public)",
        "  3. A short pack of screenshots from the iOS / Android / web app",
        "",
        "We can travel to Munich at the Consulate's convenience and would be glad "
        "to align with a community event, Open House Day or other suitable schedule.",
        "",
        "With respect and Jai Hind,",
        "",
        "Jaya Prakash Jain  \u00b7  Co-founder, IndLokal",
        "Schoeneberger Weg 14, 71065 Sindelfingen",
        "M: +49 [number]   \u00b7   E: jaya@indlokal.de",
        "",
        "Dhiraj Shah  \u00b7  Co-founder, IndLokal",
        "Bluetenaecker 17/1, 71332 Waiblingen",
        "M: +49 [number]   \u00b7   E: dhiraj@indlokal.de",
        "",
        "Web (live from May/June 2026): indlokal.de",
        "Legal entity: gemeinnuetzige UG (haftungsbeschraenkt) i.G.",
    ])

    # ------------------------------------------------------------------
    # Section 3 \u2014 Follow-up
    # ------------------------------------------------------------------
    add_h(doc, "3. Follow-up email (send after 7 working days, no reply)")

    add_p(doc, "Subject (reply-on-thread): IndLokal \u2014 gentle follow-up on our request for a 30-min introduction", italic=True)

    add_email_block(doc, [
        "Respected Sir / Madam,",
        "",
        "Following up gently on the note below from [date], in case it was missed "
        "in the inbox. We remain very keen to introduce IndLokal to the Consulate "
        "and would happily adapt to any time and format that suits the Consulate's "
        "schedule \u2014 30 minutes in person, video, or even a brief call.",
        "",
        "If there is a more appropriate point of contact for community-engagement "
        "and Indian-diaspora initiatives within the Consulate, we would be most "
        "grateful for a forwarding pointer.",
        "",
        "With renewed respect and Jai Hind,",
        "",
        "Jaya Prakash Jain  \u00b7  Dhiraj Shah",
        "Co-founders, IndLokal",
        "hello@indlokal.de  \u00b7  indlokal.de",
    ])

    # ------------------------------------------------------------------
    # Section 4 \u2014 30-min agenda
    # ------------------------------------------------------------------
    add_h(doc, "4. 30-minute meeting agenda")

    add_callout(
        doc,
        "Goal of the meeting",
        "Walk away with a verbal yes for a Letter of Support, a named point of "
        "contact at the Consulate for ongoing coordination, and \u2014 if the moment is "
        "right \u2014 a soft signal on whether ICCR / community-welfare / Pravasi "
        "Bharatiya channels are worth pursuing. Do NOT push for funding in this "
        "first meeting.",
    )

    add_h(doc, "Time-boxed structure", level=2)
    add_bullet(doc, "0\u20133 min  \u2014  Greetings, mutual introductions, founders' background (Indian citizens, long-term Germany residents, professional context).")
    add_bullet(doc, "3\u20138 min  \u2014  The problem: fragmented information for new arrivals, anchored in Destatis numbers and one personal story (e.g. how each founder found their first community event in Germany).")
    add_bullet(doc, "8\u201316 min  \u2014  Live product demo on a laptop / phone: web app + iOS app. Show: discover communities, event detail, multilingual content, submit a community, accessibility features. Keep it tight \u2014 no more than 8 minutes.")
    add_bullet(doc, "16\u201322 min  \u2014  The 12-month plan: Stuttgart anchor at soft launch; Munich beta-live months 1\u20133; Frankfurt beta-live months 3\u20136; 6\u20138 German metros by month 12. KPIs (1.5\u20132.5k MAU per anchor city, 30+ communities, 150+ resources), legal structure (gUG i.G.), funding picture (self-funded so far, 85k lean / 150k full / 250k regional grant tiers being pursued).")
    add_bullet(doc, "22\u201327 min  \u2014  The ask: (a) Letter of Support, (b) cross-listing of Consulate cultural / community events on IndLokal at no cost to the Consulate, (c) named point of contact, (d) guidance on ICCR / community-welfare channels.")
    add_bullet(doc, "27\u201330 min  \u2014  Wrap-up: confirm next concrete step, exchange business cards / WhatsApp, agree on a follow-up email summarising what was discussed.")

    # ------------------------------------------------------------------
    # Section 5 \u2014 Demo video script
    # ------------------------------------------------------------------
    add_h(doc, "5. Product demo video script (90\u2013120 seconds)")

    add_p(doc, "Format: screen recording of the iOS app (primary) intercut with web app, "
              "founder voice-over in English, soft Indian-classical / contemporary music bed at low volume, "
              "burned-in subtitles in English. File: IndLokal_Demo.mp4, max 50 MB.", italic=True)

    add_h(doc, "Scene-by-scene", level=2)
    add_bullet(doc, "0:00\u20130:08  \u2014  Title card: 'IndLokal \u2014 connecting the Indian community in Germany'. Founder names, city, date.")
    add_bullet(doc, "0:08\u20130:20  \u2014  Voice-over over a stock-style shot of a young Indian arriving at Frankfurt airport: 'When you arrive in Germany, the hardest thing is not the bureaucracy. It is finding your people \u2014 your festivals, your temple, your community.'")
    add_bullet(doc, "0:20\u20130:35  \u2014  Open the IndLokal mobile app. Show the home feed: upcoming community events in Stuttgart. Tap on a Diwali event. Show the event detail page \u2014 organiser, venue, RSVP, language tags.")
    add_bullet(doc, "0:35\u20130:50  \u2014  Switch to the Communities tab. Scroll through verified communities (temple, language circle, professional network). Tap one. Show its events, contact info, and 'verified' badge.")
    add_bullet(doc, "0:50\u20131:05  \u2014  Switch to the Resources tab. Show curated city-specific resources \u2014 Anmeldung tips, Indian grocery stores, school contacts, language schools \u2014 in English and DE, with Hindi/Tamil/Telugu translations available.")
    add_bullet(doc, "1:05\u20131:20  \u2014  Quick cut to web app at indlokal.de: same data, desktop layout, accessibility highlights (font size, screen-reader friendly).")
    add_bullet(doc, "1:20\u20131:35  \u2014  Voice-over: 'Built and funded by two India-origin founders in the Stuttgart region. Free for the community. A non-profit in formation. Public launch: May/June 2026.'")
    add_bullet(doc, "1:35\u20131:50  \u2014  Closing card: 'IndLokal \u00b7 indlokal.de \u00b7 hello@indlokal.de \u00b7 Made with respect, in Germany, for the Indian community.'")

    # ------------------------------------------------------------------
    # Section 6 \u2014 Screenshot pack
    # ------------------------------------------------------------------
    add_h(doc, "6. Screenshot pack \u2014 what to include in IndLokal_Screenshots.pdf")

    add_p(doc, "Compile into a single PDF (1 screenshot per page, light caption underneath). "
              "Aim for ~10 pages total; iOS first, web second. File: IndLokal_Screenshots.pdf, < 8 MB.", italic=True)

    add_bullet(doc, "iOS \u2014 Home / discover feed with upcoming Indian community events in Stuttgart.")
    add_bullet(doc, "iOS \u2014 Event detail page (e.g. a Diwali, Holi or Pongal event).")
    add_bullet(doc, "iOS \u2014 Communities tab with verified badges visible.")
    add_bullet(doc, "iOS \u2014 Single community detail (temple or association) with description and contact info.")
    add_bullet(doc, "iOS \u2014 Resources tab with city-specific resources (Anmeldung, schools, grocery, language).")
    add_bullet(doc, "iOS \u2014 Language switcher showing DE / EN / Hindi / Tamil / Telugu.")
    add_bullet(doc, "iOS \u2014 Submit-a-community / submit-an-event flow.")
    add_bullet(doc, "Android \u2014 Same home feed (parity proof).")
    add_bullet(doc, "Web \u2014 indlokal.de home page (desktop view).")
    add_bullet(doc, "Web \u2014 Accessibility / privacy / Impressum-style footer (proves DSGVO seriousness).")

    add_callout(
        doc,
        "Privacy hygiene before sending",
        "Screenshots and demo video MUST NOT show any real user names, emails, phone "
        "numbers, profile photos or location data. Use placeholder community names, "
        "stock-style organiser names ('Cultural Association e.V.'), and generic event "
        "titles. Confirm GDPR-clean before any external send.",
        fill="FFE8E0",
    )

    # ------------------------------------------------------------------
    # Section 7 \u2014 Briefing notes
    # ------------------------------------------------------------------
    add_h(doc, "7. Pre-meeting briefing notes \u2014 talking points and likely questions")

    add_h(doc, "Tone & posture", level=2)
    add_bullet(doc, "Respectful, formal, in English. Always 'Sir / Madam' until invited otherwise.")
    add_bullet(doc, "Lead with intent to serve the Indian community, not the funding ask.")
    add_bullet(doc, "Be transparent about pre-launch status and the entity-in-formation \u2014 honesty builds trust.")
    add_bullet(doc, "Bring physical printouts of the one-pager (3 copies) and a charged laptop + phone for the live demo.")

    add_h(doc, "Anticipated questions and rehearsed answers", level=2)
    add_bullet(doc, "Q: 'Are you a registered organisation?' \u2014 A: Yes, we are in formation as a gemeinnuetzige UG. Steuerberater is engaged; registration expected in 4\u20136 weeks. Happy to share the Notar timeline.")
    add_bullet(doc, "Q: 'Who funds you?' \u2014 A: 100% founder self-funded to date. We are now applying to AMIF (EU), ESF+ Baden-Wuerttemberg, and German foundations (Bosch, Mercator, BW Stiftung). Tier ask: 85k lean / 150k full / 250k regional.")
    add_bullet(doc, "Q: 'How is this different from existing Facebook / WhatsApp groups?' \u2014 A: Verified, multilingual, accessible, GDPR-clean, mobile-first, and integrated with official city structures. Not a chat group \u2014 a curated information layer.")
    add_bullet(doc, "Q: 'How do you keep content accurate?' \u2014 A: Mix of community-submitted content with editorial review, AI-assisted ingestion of public event sources, and a verified-organiser badge.")
    add_bullet(doc, "Q: 'Will you charge users?' \u2014 A: Never. The platform is free for the community. Long-term sustainability is via sponsored business listings and city/foundation sponsorship of the public-good content layer.")
    add_bullet(doc, "Q: 'What do you actually want from us?' \u2014 A: A short letter of support that confirms our work is relevant to the Indian community, and a named point of contact for cross-promotion of Consulate community events. Funding is a longer conversation we are happy to have separately.")
    add_bullet(doc, "Q: 'Are you affiliated politically or religiously?' \u2014 A: No. We are a non-partisan, non-sectarian platform. We list cultural events from all Indian communities \u2014 Hindu, Muslim, Sikh, Christian, Jain, Buddhist, regional and linguistic \u2014 on equal footing.")
    add_bullet(doc, "Q: 'Why should the Consulate cross-promote your platform?' \u2014 A: We extend the Consulate's reach into informal community channels at zero cost to the Consulate, with full editorial control left to the Consulate over its own listings.")

    add_h(doc, "Things to bring on the day", level=2)
    add_bullet(doc, "3 printed copies of IndLokal_Consulate_OnePager.")
    add_bullet(doc, "Laptop with the latest beta build pre-loaded; offline-capable demo path.")
    add_bullet(doc, "Both founder phones with the iOS and Android apps installed.")
    add_bullet(doc, "Business cards (or a printed visiting-card sheet if not yet printed).")
    add_bullet(doc, "Notebook for capturing the Consulate's specific asks and language for the support letter.")

    add_h(doc, "Within 24 hours after the meeting", level=2)
    add_bullet(doc, "Send a thank-you email summarising the discussion, the agreed next step, and \u2014 if a Letter of Support is on the table \u2014 a draft 1-paragraph support letter for the Consulate to adapt on its letterhead.")
    add_bullet(doc, "Add the named point of contact to a private CRM (or a tracking sheet) with date, topics discussed and next-action date.")
    add_bullet(doc, "Send a separate one-line note to the other Consulate (Frankfurt or Munich) referencing the first meeting \u2014 'we just met your colleagues in [city]; would value a similar conversation' \u2014 momentum compounds.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()

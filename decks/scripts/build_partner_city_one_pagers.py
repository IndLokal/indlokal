"""Generate IndLokal one-pagers for partner municipalities (Sindelfingen, Waiblingen).

These are scaled-down variants of the Stadt Stuttgart one-pager, framed as
"Burger:in dieser Stadt" + "Partnerkommune im Region-Stuttgart-Pilot".

Run:
    python3 decks/scripts/build_partner_city_one_pagers.py

Outputs:
    decks/output/IndLokal_Sindelfingen_OnePager.docx
    decks/output/IndLokal_Waiblingen_OnePager.docx
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUT_DIR = Path(__file__).resolve().parents[1] / "output"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT = RGBColor(0xFF, 0x8C, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# City config
# ---------------------------------------------------------------------------

@dataclass
class CityPitch:
    slug: str
    city: str
    plz: str
    sender_name: str
    sender_address: str
    sender_plz_city: str
    header_subtitle: str  # short tag in header band
    intro_paragraph: str  # lead "worum es geht" paragraph (city-specific)
    why_partner_bullets: list[tuple[str, str]]  # (bold_lead, rest)
    local_hooks_table: list[tuple[str, str]]  # 2-col mini table
    ask_intro: str
    addressee_hint: str  # used in footer/sender block


SINDELFINGEN = CityPitch(
    slug="Sindelfingen",
    city="Sindelfingen",
    plz="71065",
    sender_name="Jaya Prakash Jain",
    sender_address="Schoeneberger Weg 14",
    sender_plz_city="71065 Sindelfingen",
    header_subtitle="Partnerkommune im Region-Stuttgart-Pilot",
    intro_paragraph=(
        "Als Burger Sindelfingens und Mitgrunder von IndLokal mochte ich "
        "Ihnen ein Projekt vorstellen, das fur unsere Stadt direkt relevant "
        "ist. Sindelfingen beheimatet eine spurbare und wachsende indische "
        "Community — vor allem Ingenieurinnen und Ingenieure bei Mercedes-"
        "Benz, Bosch (Renningen) und im weiteren Umfeld, oft mit ihren "
        "Familien neu in Deutschland. Die ersten 90 Tage entscheiden uber "
        "den Integrationserfolg, doch zentrale Information uber lokale "
        "Communities, Veranstaltungen und vertrauenswurdige Anlaufstellen "
        "ist heute in geschlossenen WhatsApp-Gruppen und veralteten "
        "Facebook-Seiten verstreut. IndLokal schliesst diese Lucke — "
        "stadtspezifisch, mehrsprachig (DE/EN, Hindi/Tamil/Telugu), mobil "
        "und mit klarer Verzahnung zu offiziellen Strukturen."
    ),
    why_partner_bullets=[
        ("Lokaler Bezug: ",
         "Ich wohne und zahle Steuern in Sindelfingen. Die Plattform soll "
         "die Stadt nicht von aussen 'erreichen', sondern von innen heraus "
         "mitgestalten."),
        ("Region-Stuttgart-Pilot: ",
         "Stadt Stuttgart fuhrt den 12-Monats-Pilot. Sindelfingen wird "
         "als Partnerkommune sichtbar — mit eigenen Veranstaltungen und "
         "Inhalten, ohne stadtisches Budget zu binden."),
        ("Konkreter Mehrwert fur die Stadt: ",
         "Sichtbare Information fur die indische Community zu Anmeldung, "
         "VHS, Kitas, Schulen, Vereinen, Gesundheits- und "
         "Beratungsangeboten in Sindelfingen — gepflegt von uns, "
         "verlinkt auf stadtische Quellen."),
        ("Belastbare Zahlen: ",
         "Quartalsweise Wirkungsberichte und unabhangige Evaluation "
         "(freier Evaluations-Consultant oder Lehrstuhl an Uni Stuttgart / "
         "Hochschule Esslingen; ifo / BAMF-FZ als Forschungs-Anker im "
         "Aufbau) liefern der Stadt erstmals belastbare Daten zur "
         "indischen Community in Sindelfingen."),
    ],
    local_hooks_table=[
        ("Mercedes-Benz Werk Sindelfingen",
         "Grosste Konzentration indischer Ingenieur:innen und Familien"),
        ("Bosch Forschungscampus Renningen",
         "Indische Forscher:innen und Entwickler:innen im direkten Umfeld"),
        ("VHS Sindelfingen + Stadtbibliothek",
         "Naturliche Verlinkungs- und Veranstaltungspartner"),
        ("Burgerstiftung Sindelfingen",
         "Potenzieller lokaler Co-Forderer fur kleine Bausteine"),
    ],
    ask_intro=(
        "Ein 30-minutiges Kennenlerngesprach mit der/dem "
        "Integrationsbeauftragten der Stadt Sindelfingen — um "
        "vorzustellen:"
    ),
    addressee_hint="Stadt Sindelfingen — Amt fur Bildung, Soziales und Kultur, Integrationsbeauftragte/r",
)


WAIBLINGEN = CityPitch(
    slug="Waiblingen",
    city="Waiblingen",
    plz="71332",
    sender_name="Dhiraj Shah",
    sender_address="Bluetenaecker 17/1",
    sender_plz_city="71332 Waiblingen",
    header_subtitle="Partnerkommune im Region-Stuttgart-Pilot",
    intro_paragraph=(
        "Als Burger Waiblingens und Mitgrunder von IndLokal mochte ich "
        "Ihnen ein Projekt vorstellen, das fur unsere Stadt und den "
        "gesamten Rems-Murr-Kreis direkt relevant ist. Mit Stihl als "
        "Ankerarbeitgeber und einem wachsenden IT- und Engineering-"
        "Korridor lebt im Rems-Murr-Kreis eine spurbar wachsende indische "
        "Community — viele neu in Deutschland, oft mit Familie. Die "
        "ersten 90 Tage entscheiden uber den Integrationserfolg, doch "
        "zentrale Information uber lokale Communities, Veranstaltungen "
        "und vertrauenswurdige Anlaufstellen ist heute in geschlossenen "
        "WhatsApp-Gruppen und veralteten Facebook-Seiten verstreut. "
        "IndLokal schliesst diese Lucke — stadtspezifisch, mehrsprachig "
        "(DE/EN, Hindi/Tamil/Telugu), mobil und mit klarer Verzahnung zu "
        "offiziellen Strukturen."
    ),
    why_partner_bullets=[
        ("Lokaler Bezug: ",
         "Ich wohne und zahle Steuern in Waiblingen. Die Plattform soll "
         "die Stadt nicht von aussen 'erreichen', sondern von innen "
         "heraus mitgestalten."),
        ("Region-Stuttgart-Pilot: ",
         "Stadt Stuttgart fuhrt den 12-Monats-Pilot. Waiblingen wird als "
         "Partnerkommune sichtbar — mit eigenen Veranstaltungen und "
         "Inhalten, ohne stadtisches Budget zu binden."),
        ("Konkreter Mehrwert fur die Stadt: ",
         "Sichtbare Information fur die indische Community zu Anmeldung, "
         "VHS, Kitas, Schulen, Vereinen, Gesundheits- und "
         "Beratungsangeboten in Waiblingen — gepflegt von uns, verlinkt "
         "auf stadtische Quellen."),
        ("Belastbare Zahlen: ",
         "Quartalsweise Wirkungsberichte und unabhangige Evaluation "
         "(freier Evaluations-Consultant oder Lehrstuhl an Uni Stuttgart / "
         "Hochschule Esslingen; ifo / BAMF-FZ als Forschungs-Anker im "
         "Aufbau) liefern der Stadt erstmals belastbare Daten zur "
         "indischen Community in Waiblingen und im Rems-Murr-Kreis."),
    ],
    local_hooks_table=[
        ("Stihl HQ + Engineering-Korridor",
         "Wachsende indische IT- und Engineering-Cohorte"),
        ("Landratsamt Rems-Murr-Kreis — Stabsstelle Integration",
         "Naturlicher Co-Adressat fur Wirkung uber Waiblingen hinaus"),
        ("VHS Unteres Remstal + Stadtbibliothek",
         "Verlinkungs- und Veranstaltungspartner"),
        ("Burgerstiftung Waiblingen",
         "Potenzieller lokaler Co-Forderer fur kleine Bausteine"),
    ],
    ask_intro=(
        "Ein 30-minutiges Kennenlerngesprach mit der/dem "
        "Integrationsbeauftragten der Stadt Waiblingen (gerne auch "
        "gemeinsam mit dem Landratsamt Rems-Murr-Kreis) — um "
        "vorzustellen:"
    ),
    addressee_hint="Stadt Waiblingen — Fachbereich Bildung und Soziales, Integrationsbeauftragte/r",
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_run(para, text: str, *, bold: bool = False, italic: bool = False,
            size: int = 10, color: RGBColor = DARK):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def tight(para, *, before: int = 0, after: int = 2, line: float = 1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    tight(p, before=4, after=2)
    add_run(p, text.upper(), bold=True, size=10, color=PRIMARY)


def bullet(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    tight(p, after=1, line=1.1)
    if bold_lead:
        add_run(p, bold_lead, bold=True, size=10)
        add_run(p, text, size=10)
    else:
        add_run(p, text, size=10)


# ---------------------------------------------------------------------------
# Build per city
# ---------------------------------------------------------------------------

def build_one(city: CityPitch) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ---- Header band ---------------------------------------------------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)
    left.width = Cm(11)
    right.width = Cm(7)

    set_cell_shading(left, "0B3D91")
    set_cell_shading(right, "FF8C00")

    p = left.paragraphs[0]
    tight(p, after=0)
    add_run(p, "IndLokal", bold=True, size=20, color=WHITE)
    p2 = left.add_paragraph()
    tight(p2, after=0)
    add_run(p2, "Integration durch Information",
            italic=True, size=11, color=WHITE)
    p3 = left.add_paragraph()
    tight(p3, after=0)
    add_run(p3,
            f"Eine digitale Integrationsinfrastruktur fur die indische "
            f"Community in {city.city}",
            size=10, color=WHITE)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp, after=0)
    add_run(rp, city.header_subtitle, bold=True, size=11, color=WHITE)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp2, after=0)
    add_run(rp2, "Lead: Stadt Stuttgart  ·  12 Monate",
            size=10, color=WHITE)
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp3, after=0)
    add_run(rp3, "Nicht verwassernd  ·  Unabhangige Evaluation",
            italic=True, size=9, color=WHITE)

    spacer = doc.add_paragraph()
    tight(spacer, after=0)

    # ---- Worum es geht -------------------------------------------------
    section_heading(doc, "Worum es geht")
    p = doc.add_paragraph()
    tight(p, after=3)
    add_run(p, city.intro_paragraph, size=10)

    # ---- Was bereits existiert ----------------------------------------
    section_heading(doc, "Was bereits existiert")
    bullet(doc, "Web-App indlokal.de und native iOS/Android-Apps live im Store.",
           bold_lead="Plattform: ")
    bullet(doc, "Verifizierte Communities, Veranstaltungen und Ressourcen — mehrsprachig vorbereitet (DE/EN, Hindi/Tamil/Telugu).",
           bold_lead="Inhalt: ")
    bullet(doc, "Komplett aus Eigenmitteln gebaut. Eine Forderung finanziert NICHT die Entwicklung, sondern Community-Arbeit, Partnerschaften und Evaluation.",
           bold_lead="Finanzierung bisher: ")

    # ---- Lokale Anknupfungspunkte ------------------------------------
    section_heading(doc, f"Lokale Anknupfungspunkte in {city.city}")
    table = doc.add_table(rows=len(city.local_hooks_table), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, body) in enumerate(city.local_hooks_table):
        row = table.rows[i].cells
        row[0].text = ""
        row[1].text = ""
        p1 = row[0].paragraphs[0]
        add_run(p1, label, bold=True, size=10, color=PRIMARY)
        p2 = row[1].paragraphs[0]
        add_run(p2, body, size=10)

    # ---- Warum wir die Stadt suchen -----------------------------------
    section_heading(doc, f"Warum wir die Stadt {city.city} als Partner suchen")
    for lead, rest in city.why_partner_bullets:
        bullet(doc, rest, bold_lead=lead)

    # ---- Die Bitte ----------------------------------------------------
    section_heading(doc, "Was wir konkret erbitten")
    p = doc.add_paragraph()
    tight(p, after=2)
    add_run(p, city.ask_intro, size=10)
    bullet(doc, "die Plattform live (Web + Mobile),")
    bullet(doc, "den geplanten Pilot-Umfang und die KPIs,")
    bullet(doc,
           f"wie eine schlanke Partnerschaft aussehen kann — "
           f"Letter of Support, ggf. kleiner Baustein aus dem "
           f"{city.city}er Integrationsfonds, gemeinsame Sichtbarkeit, "
           f"Verlinkung mit stadtischen Angeboten.")

    # ---- Sender / contact band ---------------------------------------
    spacer = doc.add_paragraph()
    tight(spacer, after=0)

    contact_table = doc.add_table(rows=1, cols=2)
    contact_table.autofit = False
    lc = contact_table.cell(0, 0)
    rc = contact_table.cell(0, 1)
    lc.width = Cm(10)
    rc.width = Cm(8)
    set_cell_shading(lc, "F2F4F8")
    set_cell_shading(rc, "F2F4F8")

    lp = lc.paragraphs[0]
    tight(lp, after=0)
    add_run(lp, "Absender", bold=True, size=10, color=PRIMARY)
    lp2 = lc.add_paragraph(); tight(lp2, after=0)
    add_run(lp2, city.sender_name, bold=True, size=10)
    lp3 = lc.add_paragraph(); tight(lp3, after=0)
    add_run(lp3, city.sender_address, size=10)
    lp4 = lc.add_paragraph(); tight(lp4, after=0)
    add_run(lp4, city.sender_plz_city, size=10)
    lp5 = lc.add_paragraph(); tight(lp5, after=0)
    add_run(lp5, "Mitgrunder, IndLokal", italic=True, size=9, color=MUTED)

    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp, after=0)
    add_run(rp, "Kontakt", bold=True, size=10, color=PRIMARY)
    rp2 = rc.add_paragraph(); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; tight(rp2, after=0)
    add_run(rp2, "hello@indlokal.de", bold=True, size=10)
    rp3 = rc.add_paragraph(); rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT; tight(rp3, after=0)
    add_run(rp3, "indlokal.de  ·  @indlokal", size=10)
    rp4 = rc.add_paragraph(); rp4.alignment = WD_ALIGN_PARAGRAPH.RIGHT; tight(rp4, after=0)
    add_run(rp4, f"Adressat: {city.addressee_hint}",
            italic=True, size=8, color=MUTED)

    out = OUT_DIR / f"IndLokal_{city.slug}_OnePager.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build() -> None:
    for city in (SINDELFINGEN, WAIBLINGEN):
        path = build_one(city)
        print(f"Wrote {path}")


if __name__ == "__main__":
    build()

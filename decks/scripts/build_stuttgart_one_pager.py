"""Generate a one-page IndLokal x Stadt Stuttgart introduction (DE/EN).

Run:
    python3 decks/scripts/build_stuttgart_one_pager.py

Output:
    decks/output/IndLokal_Stuttgart_OnePager.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "IndLokal_Stuttgart_OnePager.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT = RGBColor(0xFF, 0x8C, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_run(para, text: str, *, bold: bool = False, italic: bool = False,
            size: int = 10, color: RGBColor = DARK):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def tight(para, *, before: int = 0, after: int = 2, line: float = 1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    tight(p, before=4, after=2)
    add_run(p, text.upper(), bold=True, size=10, color=PRIMARY)


def bullet(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    tight(p, after=1, line=1.1)
    if bold_lead:
        add_run(p, bold_lead, bold=True, size=10)
        add_run(p, text, size=10)
    else:
        add_run(p, text, size=10)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()

    # tighten page to fit one page
    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ---- Header band ---------------------------------------------------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)
    left.width = Cm(11)
    right.width = Cm(7)

    set_cell_shading(left, "0B3D91")
    set_cell_shading(right, "FF8C00")

    p = left.paragraphs[0]
    tight(p, after=0)
    add_run(p, "IndLokal", bold=True, size=20,
            color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = left.add_paragraph()
    tight(p2, after=0)
    add_run(p2, "Integration durch Information",
            italic=True, size=11, color=RGBColor(0xFF, 0xFF, 0xFF))
    p3 = left.add_paragraph()
    tight(p3, after=0)
    add_run(p3,
            "Eine digitale Integrationsinfrastruktur fur die indische Community in Stuttgart",
            size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp, after=0)
    add_run(rp, "12-Monats-Pilot", bold=True, size=11,
            color=RGBColor(0xFF, 0xFF, 0xFF))
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp2, after=0)
    add_run(rp2, "Stuttgart  \u00b7  EUR 85k lean / 150k voll", size=10,
            color=RGBColor(0xFF, 0xFF, 0xFF))
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(rp3, after=0)
    add_run(rp3, "Stack-fahig  \u00b7  Mit unabhangiger Evaluation",
            italic=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))

    spacer = doc.add_paragraph()
    tight(spacer, after=0)

    # ---- The opportunity ----------------------------------------------
    section_heading(doc, "Worum es geht")
    p = doc.add_paragraph()
    tight(p, after=3)
    add_run(p,
            "Rund 17.000 Inderinnen und Inder leben heute in Stuttgart "
            "(Stadt Stuttgart 2024) — Tendenz stark steigend. Bundesweit "
            "ist es die am schnellsten wachsende Drittstaaten-Migration "
            "(+22% YoY, Destatis). Doch zentrale Information uber lokale "
            "Communities, kulturelle Veranstaltungen und vertrauenswurdige "
            "Anlaufstellen ist fragmentiert: geschlossene WhatsApp-Gruppen, "
            "veraltete Facebook-Seiten, Mundpropaganda. Die ersten 90 Tage "
            "in Deutschland entscheiden uber Integrationserfolg — und genau "
            "hier fehlt eine kuratierte, mehrsprachige Schicht.",
            size=10)

    p = doc.add_paragraph()
    tight(p, after=3)
    add_run(p, "IndLokal schliesst diese Lucke. ", bold=True, size=10)
    add_run(p,
            "Verifizierte Communities, frische Veranstaltungen und "
            "stadtspezifische Ressourcen — mobil, mehrsprachig, "
            "datensparsam, mit klarer Verbindung zu offiziellen Strukturen "
            "wie dem Welcome Center und der Volkshochschule.",
            size=10)

    # ---- What is built today ------------------------------------------
    section_heading(doc, "Stand heute (ehrlich)")
    bullet(doc, "Web-App indlokal.de und native iOS/Android-Apps fertig entwickelt, in privatem Beta. Soft-Launch geplant in den naechsten ~4 Wochen.",
           bold_lead="Plattform: ")
    bullet(doc, "Verifizierte Communities, Veranstaltungen und Ressourcen \u2014 mehrsprachig vorbereitet (DE/EN, Hindi/Tamil/Telugu).",
           bold_lead="Inhalt: ")
    bullet(doc, "Komplett aus Eigenmitteln der zwei Gruender entwickelt. Kein externes Kapital, noch keine Umsaetze. Eine Foerderung finanziert NICHT die Entwicklung, sondern Launch, Community-Arbeit, Partnerschaften und Evaluation.",
           bold_lead="Finanzierung bisher: ")
    bullet(doc, "Gemeinnuetzige UG / gGmbH in Gruendung \u2014 Abschluss innerhalb des Foerderentscheidungs-Zeitraums vorgesehen.",
           bold_lead="Rechtsform: ")

    # ---- The pilot ----------------------------------------------------
    section_heading(doc, "Der Stuttgart-Pilot — 12 Monate, konkrete Ergebnisse")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Reichweite", "Communities", "Ressourcen", "Partnerschaften"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        pp = c.paragraphs[0]
        add_run(pp, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
    row = table.add_row().cells
    values = [
        "1.500\u20132.500 aktive Nutzer:innen / Monat (10\u201315% der indischen Community in Stuttgart); Stretch-Ziel 4.000",
        "30+ verifizierte Vereine, Tempel, Sprachkreise, Berufsnetzwerke (~40 Kandidaten heute identifiziert)",
        "150+ stadtspezifische Inhalte; 60+ professionell ubersetzt (Hindi/Tamil/Telugu)",
        "Letters of Support von Welcome Center, IHK Stuttgart, mind. 1 Hochschule",
    ]
    for i, v in enumerate(values):
        row[i].text = ""
        pp = row[i].paragraphs[0]
        add_run(pp, v, size=9)

    # ---- Why Stuttgart Welcome Center is critical ---------------------
    section_heading(doc, "Warum wir die Stadt Stuttgart als Partner suchen")
    bullet(doc,
           "Sichtbarkeit & Vertrauen: ein Brief der Abteilung Integration ist fur AMIF, ESF+ BW, BW Stiftung, Mercator und Bosch faktisch Voraussetzung.",
           bold_lead="Letter of Support: ")
    bullet(doc,
           "Wechselseitige Verlinkung mit Welcome Center / Integrationsbeauftragte — keine Doppelstrukturen, sondern gegenseitige Verstarkung.",
           bold_lead="Verzahnung: ")
    bullet(doc,
           "Open-Data-Schnittstelle fur stadtische Integrationsangebote, Sprachkurse und Veranstaltungen — die Stadt liefert Daten, IndLokal verteilt sie zielgruppengerecht.",
           bold_lead="Daten-Partnerschaft: ")
    bullet(doc,
           "Quartalsweise Wirkungsberichte und unabhangige Evaluation (freier Evaluations-Consultant oder Lehrstuhl Uni Stuttgart / HS Esslingen; ifo / BAMF-FZ als Forschungs-Anker im Aufbau) \u2014 die Stadt erhalt erstmals belastbare Daten zur indischen Community in Stuttgart.",
           bold_lead="Evidenz: ")

    # ---- The ask ------------------------------------------------------
    section_heading(doc, "Was wir konkret erbitten")
    p = doc.add_paragraph()
    tight(p, after=2)
    add_run(p, "Ein 30-minutiges Kennenlerngesprach ", bold=True, size=10)
    add_run(p, "mit der Abteilung Integration der Stadt Stuttgart, um vorzustellen:",
            size=10)
    bullet(doc, "die Plattform live (Web + Mobile),")
    bullet(doc, "den geplanten Pilot-Umfang und KPIs,")
    bullet(doc, "wie eine schlanke Partnerschaft (Letter of Support, ggf. Daten-Schnittstelle, gemeinsame Sichtbarkeit) aussehen kann — ohne stadtisches Budget zu binden.")

    # ---- Footer / contact --------------------------------------------
    spacer = doc.add_paragraph()
    tight(spacer, after=0)

    contact_table = doc.add_table(rows=1, cols=1)
    contact_cell = contact_table.cell(0, 0)
    set_cell_shading(contact_cell, "F2F4F8")
    contact_cell.paragraphs[0].text = ""
    cp = contact_cell.paragraphs[0]
    tight(cp, after=0)
    add_run(cp, "Kontakt  ", bold=True, size=10, color=PRIMARY)
    add_run(cp,
            "Grundungsteam IndLokal  ·  Stuttgart, Baden-Wurttemberg  ·  ",
            size=10)
    add_run(cp, "hello@indlokal.de", bold=True, size=10, color=PRIMARY)
    add_run(cp, "  ·  indlokal.de  ·  @indlokal", size=10)

    # ---- Bilingual line ----------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(p, before=4, after=0)
    add_run(p,
            "Dieses Dokument liegt auch auf Englisch vor. / "
            "An English version of this one-pager is available on request.",
            italic=True, size=8, color=MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()

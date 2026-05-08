"""Generate the IndLokal grant action plan: what to do with the existing docs.

Run:
    python3 decks/scripts/build_grant_action_plan.py

Output:
    decks/output/IndLokal_Grant_Action_Plan.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "IndLokal_Grant_Action_Plan.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT = RGBColor(0xFF, 0x8C, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GOOD = RGBColor(0x1B, 0x7A, 0x3E)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [(1, 20), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = PRIMARY
        s.font.bold = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = MUTED
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = rule.add_run("— ★ —")
    r3.font.color.rgb = ACCENT
    r3.font.size = Pt(13)
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_checklist(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph()
        r = p.add_run("☐  ")
        r.font.size = Pt(12)
        r.bold = True
        p.add_run(it).font.size = Pt(11)


def add_callout(doc: Document, title: str, body: str, fill: str = "FFF8DC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}\n")
    r.bold = True
    r.font.color.rgb = PRIMARY
    r.font.size = Pt(12)
    p.add_run(body).font.size = Pt(11)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    set_base_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(
        doc,
        "IndLokal — Grant Action Plan",
        "What to do with the docs we have, in order",
    )
    # --- Reality check up top
    add_callout(
        doc,
        "Status today (be honest with yourselves)",
        "Platform built, in private beta. Soft public launch planned in "
        "~4 weeks. No legal entity yet (still deciding between gUG, "
        "e.V., gGmbH). No external capital, no revenue. This sequence "
        "is built around that reality \u2014 not around pretending we are "
        "further along than we are.",
        fill="FFE8E0",
    )

    # --- Phase 0: prerequisites
    add_h(doc, "Phase 0 \u2014 weeks 1\u20134: prerequisites (do BEFORE any funder ask)", 1)
    add_p(
        doc,
        "Without these three things, no foundation will write a cheque "
        "and no public funder will accept the application. Sequence them "
        "in parallel; they all complete in roughly the same window.",
    )

    add_h(doc, "P0.1 \u2014 Legal entity decision and registration", 2)
    add_table(
        doc,
        ["Option", "Min. capital", "People", "Pros", "Cons"],
        [
            ["gUG (gemeinn\u00fctzig)", "\u20ac1", "1\u20132 founders",
             "Cheap, fast (~3\u20136 weeks), unlocks Google/Microsoft for Nonprofits, can convert to gGmbH later",
             "Needs Steuerberater + Notar; Gemeinn\u00fctzigkeit must be confirmed by Finanzamt"],
            ["e.V.", "\u20ac0", "7 founding members min.",
             "Free, classic civil-society form, strong with Land BW funders",
             "7 members hard for a 2-founder team; less corporate-friendly"],
            ["gGmbH", "\u20ac25k", "1\u20132 founders",
             "Most credible, full liability cover, foundation-friendly",
             "\u20ac25k tied capital; slower"],
        ],
    )
    add_callout(
        doc,
        "Recommended path",
        "For a 2-founder pre-launch project: start as gemeinn\u00fctzige UG "
        "(haftungsbeschr\u00e4nkt), convert to gGmbH later when capital and "
        "track record exist. Talk to a Steuerberater specialised in "
        "Gemeinn\u00fctzigkeit this week. Realistic closing: 4\u20136 weeks.",
    )
    add_checklist(
        doc,
        [
            "Pick 2 Steuerberater specialised in Gemeinn\u00fctzigkeit, get fee quotes (typically \u20ac1.5\u20133k)",
            "Decide entity type (recommend gUG)",
            "Draft Satzung with Steuerberater (must hit \u00a7 52 AO purposes: Integration, V\u00f6lkerverst\u00e4ndigung, Bildung)",
            "Notar-Termin for incorporation",
            "Open business bank account",
            "Submit to Finanzamt for Gemeinn\u00fctzigkeitsbescheinigung",
        ],
    )

    add_h(doc, "P0.2 \u2014 Soft launch the platform (~4 weeks)", 2)
    add_checklist(
        doc,
        [
            "Lock launch date (target: ~4 weeks from today)",
            "Seed 10\u201315 communities and 50\u2013100 resources before launch \u2014 'live in beta with real content' beats 'will launch when funded'",
            "Publish accessibility statement (BFSG)",
            "Publish DSGVO/Impressum legally clean",
            "Run 1 small real community event through the platform; capture RSVPs and attendance",
            "Capture 5 user quotes from new arrivals \u2014 strongest evidence in any application",
        ],
    )

    add_h(doc, "P0.3 \u2014 Reality-check sprint (1\u20132 days)", 2)
    add_checklist(
        doc,
        [
            "Build 'real community list' spreadsheet: 30+ Indian organisations across Stuttgart, Sindelfingen, Waiblingen with name, contact, status",
            "Talk to 5 Indians who arrived in the last 12 months. Verbatim quotes go into pitch materials",
            "Email 1 freelance evaluation consultant: 'capacity for a 12mo integration evaluation, \u20ac15k?'",
            "Light-touch email to ifo or BAMF-FZ researcher: 20-min advice call (not a contract ask)",
        ],
    )

    add_callout(
        doc,
        "Why Phase 0 is non-negotiable",
        "Without legal entity, AMIF / ESF+ / Mercator / Bosch / BW Stiftung "
        "all reject the application at the door. Without a live product "
        "with even 10 communities, the strongest deck claim ('platform "
        "built') reads as vapourware. Spend 4 weeks here and the next 12 "
        "months become 10\u00d7 easier. Skip it and you waste 6 months on "
        "applications no one can accept.",
        fill="E8F0FF",
    )
    # --- 0. The artefacts we have
    add_h(doc, "0. What we have in hand", 1)
    add_table(
        doc,
        ["File", "Purpose", "Send to"],
        [
            ["indlokal-grant-deck.pptx",
             "Full pitch deck (3 funding tiers, KPIs, budget, timeline)",
             "Foundations + AMIF/ESF+ programme officers, after first contact"],
            ["IndLokal_Stuttgart_OnePager.docx",
             "1-page intro for Stadt Stuttgart Abteilung Integration",
             "Stadt Stuttgart Welcome Center / Integrationsbeauftragte"],
            ["IndLokal_Sindelfingen_OnePager.docx",
             "1-page intro under Jaya's name, Sindelfingen-resident framing",
             "Stadt Sindelfingen Integrationsbeauftragte"],
            ["IndLokal_Waiblingen_OnePager.docx",
             "1-page intro under Dhiraj's name, Waiblingen-resident framing",
             "Stadt Waiblingen + Landratsamt Rems-Murr-Kreis Integration"],
            ["IndLokal_Grant_Strategy.docx",
             "Internal strategy memo (programme tiers, gGmbH/e.V., paperwork)",
             "Internal use only — not for funders"],
            ["EU_Grants_Explained.docx",
             "Background primer on EU grant landscape",
             "Internal reference / orientation for new team members"],
        ],
    )
    add_callout(
        doc,
        "Golden rule",
        "Never send the deck cold. Always send the relevant 1-pager first, "
        "ask for a 30-min meeting, and bring the deck to the meeting. "
        "1-pagers get read; cold decks get archived.",
    )

    # --- 1. Reality-check sprint
    add_h(doc, "1. Week 1 — reality-check sprint (do BEFORE any funder contact)", 1)
    add_p(
        doc,
        "These five tasks calibrate the numbers in the deck. Without them, "
        "any funder conversation is built on sand. Allocate 1\\u20132 days.",
    )
    add_checklist(
        doc,
        [
            "Build the 'real community list' spreadsheet: every Indian organisation, temple, association, sports club, language circle, alumni network you can name in Stuttgart / Sindelfingen / Waiblingen. Aim for 30+. Columns: name, contact person, contact channel, status (cold / warm / committed).",
            "Talk to 5 Indians who arrived in the region in the last 12 months. One question: 'Where did you find your first community event?' Capture verbatim quotes — these become the deck's strongest evidence.",
            "Run one real event through the platform: a community sponsor posts, you measure RSVPs and attendance. Even 12 RSVPs gives you 'baseline' data instead of 'baseline 0'.",
            "Email 1 freelance evaluation consultant (or a chair at Uni Stuttgart / HS Esslingen) and ask: 'capacity to evaluate a 12mo integration pilot for €15k?' Their reply unlocks the evaluation line item.",
            "Email 1 BAMF-FZ or ifo researcher (light touch): 'we'd love a 20-min call to learn how you'd advise on evaluation design'. NOT a contract ask. Builds the research-link credibility for later.",
        ],
    )

    # --- 2. Anchor meeting
    add_h(doc, "Phase 1 \u2014 weeks 4\u20136: anchor meetings, in PARALLEL", 1)
    add_callout(
        doc,
        "Run city + Consulate outreach in parallel, not sequentially",
        "In the AI era, 'pilot one city, prove it, then expand' is too slow. "
        "Per-city marginal cost is low; trust and partnerships are the real "
        "bottleneck. So the anchor-meeting wave runs across Stuttgart (anchor city), "
        "Munich and Frankfurt (city governments + Indian Consulates General) "
        "simultaneously. Use IndLokal_Consulate_Outreach_Pack.docx for the "
        "diplomatic track and the city one-pagers for the municipal track.",
        fill="E8F0FF",
    )
    add_callout(
        doc,
        "Single most important action",
        "Book a 30-minute meeting with Stadt Stuttgart Abteilung "
        "Integration. Their Letter of Support unlocks AMIF, ESF+ BW, "
        "BW Stiftung, Bosch and Mercator. One meeting = five "
        "applications enabled.",
        fill="E8F0FF",
    )
    add_numbered(
        doc,
        [
            "Send IndLokal_Stuttgart_OnePager.docx to the Abteilung Integration mailbox (Stadt Stuttgart, Eberhardstraße). Use the email template in section 8 below.",
            "If no reply in 7 days, follow up once. If still no reply, ask a known intermediary (IHK contact, university international office, community leader) for a warm intro.",
            "In parallel, both founders book meetings in their home cities (Sindelfingen / Waiblingen) using their respective one-pagers. Don't wait for Stuttgart.",
            "Goal of each meeting: a verbal yes for a Letter of Support. Bring the deck on a laptop. Show the live app. Don't ask for money in this meeting.",
        ],
    )

    # --- 3. Paperwork in parallel
    add_h(doc, "Phase 2 \u2014 weeks 4\u20138: grant-paperwork basics (after entity exists)", 1)
    add_p(
        doc,
        "These items have long lead times. Start them on day 1 even "
        "though you won't use them for weeks.",
    )
    add_checklist(
        doc,
        [
            "Register PIC code on the EU Funding & Tenders Portal (free, slow). Required for AMIF.",
            "Both founders create EU Login accounts.",
            "Talk to a Steuerberater specialised in Gemeinnützigkeit about gGmbH or e.V. setup. Get written cost + timeline. Decision: yes / no / later.",
            "Draft DSGVO concept: DPIA for community/event data, AVV with hosting providers, ROPA (Verarbeitungsverzeichnis).",
            "Publish accessibility statement on indlokal.de per BFSG.",
            "Open a separate project bank account (will be required by AMIF / ESF+ if won).",
            "Pull together last 1\\u20132 years of financials, audited or reviewed.",
        ],
    )

    # --- 4. First applications (foundations)
    add_h(doc, "Phase 3 \u2014 weeks 6\u201312: first cheque attempts (foundations)", 1)
    add_p(
        doc,
        "Foundations have rolling deadlines and faster decisions than "
        "public funding. Aim to land at least one small cheque before "
        "tackling AMIF / ESF+.",
    )
    add_table(
        doc,
        ["Funder", "Tier from deck", "Application channel", "Timeline"],
        [
            ["Stadt Stuttgart Integrationsfonds", "€85k lean (top of range or contribution)",
             "Online form via stadt-stuttgart.de + cover letter + deck",
             "Decision in 4\\u20138 weeks"],
            ["Bürgerstiftung Stuttgart", "Small contribution",
             "Online form + 2-page concept",
             "Quarterly decisions"],
            ["Robert Bosch Stiftung — Migration", "€50\\u201385k",
             "Email programme officer + deck. Stuttgart-HQ helps.",
             "Decision in 8\\u201316 weeks"],
            ["Stiftung Mercator — Teilhabe & Zusammenhalt", "€85\\u2013150k",
             "Email programme officer + deck",
             "Decision in 12\\u201320 weeks"],
            ["BW Stiftung — Integration", "€85\\u2013150k",
             "Public calls (check site) or unsolicited concept paper",
             "Call-cycle dependent"],
        ],
    )
    add_callout(
        doc,
        "Stack-target reminder",
        "Each foundation is a candidate for the €85k lean OR a slice of "
        "the €150k full ask. Keep a tracking sheet of who you've asked "
        "for what. Don't double-promise the same euros.",
    )

    # --- 5. AMIF / ESF+
    add_h(doc, "Phase 4 \u2014 weeks 10\u201320: public funding (AMIF / ESF+)", 1)
    add_p(
        doc,
        "Larger cheques, longer timelines, more paperwork. Start once "
        "you have at least one Letter of Support and ideally one "
        "foundation win or commitment.",
    )
    add_numbered(
        doc,
        [
            "Email BAMF Regionalkoordinator BW: introduce IndLokal, ask about current AMIF national programme call schedule. Attach the deck.",
            "Email BW Ministerium für Soziales referat for Integrationsförderung. Attach deck. Ask which open call best fits.",
            "Identify the next AMIF or ESF+ BW open call. Confirm eligibility, deadlines, co-financing requirement.",
            "Draft Part A (administrative) and Part B (technical proposal). Allocate 80\\u2013120 hours.",
            "Get NCP (national contact point) feedback on the draft. Free, valuable.",
            "Submit at least 24h before deadline.",
        ],
    )

    # --- 6. Corporate sponsorship
    add_h(doc, "Phase 5 \u2014 months 2\u20136: corporate sponsorship (parallel track)", 1)
    add_p(
        doc,
        "Often faster than grants and counts as co-financing for "
        "AMIF / ESF+. Target Indian-major and German companies with "
        "Indian workforce.",
    )
    add_bullets(
        doc,
        [
            "Indian IT majors with Stuttgart-region offices: TCS, Infosys, Wipro, HCL, Tech Mahindra. Approach via diaspora-affairs / CSR contact.",
            "German employers with large Indian workforce: Mercedes-Benz (Sindelfingen), Bosch (Renningen / Stuttgart), Porsche (Weissach), SAP (Walldorf), Allianz (Stuttgart).",
            "Indo-German Chamber of Commerce (IGCC) — they can intro to multiple corporates and amplify via newsletter.",
            "Ask for: €5\\u201325k sponsorship, optional in-kind (event space, employee volunteer hours, content channels).",
        ],
    )

    # --- 7. In-kind & advertising channels
    add_h(doc, "Phase 6 \u2014 once the gemeinn\u00fctzig entity is registered", 1)
    add_checklist(
        doc,
        [
            "Apply for Google for Nonprofits → unlocks Ad Grants (~€9k/month) + Workspace.",
            "Apply for Microsoft for Nonprofits → Azure credits (potentially replaces hosting line item).",
            "Apply for Apple App Store nonprofit fee waiver.",
            "Apply for GitHub for Nonprofits / Notion / Slack / Figma nonprofit programmes.",
            "Apply to Prototype Fund (BMBF + OKF DE) if you can release one component as OSS.",
        ],
    )

    # --- 8. Email templates
    add_h(doc, "Email templates", 1)

    add_h(doc, "Cold intro to a city Integrationsbeauftragte (DE)", 2)
    add_callout(
        doc,
        "Subject: IndLokal — digitale Integrationsinfrastruktur für die indische Community in [Stadt]",
        "Sehr geehrte Frau / Herr [Name],\n\n"
        "mein Name ist [Vorname Nachname], ich bin Bürger:in [Stadt] und "
        "Mitgründer:in von IndLokal — einer digitalen Plattform, die "
        "neuen indischen Zuwander:innen hilft, lokale Communities, "
        "Veranstaltungen und vertrauenswürdige Anlaufstellen in ihrer "
        "neuen Heimatstadt zu finden.\n\n"
        "Die Plattform ist live (indlokal.de + iOS/Android), aus "
        "Eigenmitteln gebaut und steht zur Verfügung. Wir bauen gerade "
        "einen 12-monatigen Stuttgart-Pilot mit Partnerkommunen "
        "Sindelfingen und Waiblingen auf und würden uns sehr über ein "
        "30-minütiges Kennenlerngespräch freuen, um:\n"
        "  - die Plattform live vorzustellen,\n"
        "  - Synergien mit dem Welcome Center / Integrationsangeboten "
        "der Stadt zu besprechen,\n"
        "  - eine schlanke Form der Zusammenarbeit zu skizzieren — "
        "ohne städtisches Budget zu binden.\n\n"
        "Im Anhang finden Sie eine einseitige Übersicht. Gerne stelle "
        "ich auch kurz das Projekt direkt im Termin vor.\n\n"
        "Herzliche Grüße\n"
        "[Vorname Nachname]\n"
        "Mitgründer:in IndLokal\n"
        "[Adresse / PLZ Stadt]\n"
        "hello@indlokal.de · indlokal.de",
        fill="F2F4F8",
    )

    add_h(doc, "Foundation programme officer intro (DE)", 2)
    add_callout(
        doc,
        "Subject: IndLokal — digitale Integrationsinfrastruktur für die indische Community in der Region Stuttgart",
        "Sehr geehrte Frau / Herr [Name],\n\n"
        "wir bauen IndLokal auf, eine digitale Integrationsinfrastruktur "
        "für die indische Community in Deutschland — beginnend mit einem "
        "12-monatigen Stuttgart-Pilot in Partnerschaft mit Stadt "
        "Stuttgart, Sindelfingen und Waiblingen.\n\n"
        "Warum wir uns an Sie wenden: Ihr Förderprogramm "
        "[Programmname / Schwerpunkt] passt direkt zu unserem Vorhaben "
        "— [1\\u20132 Sätze, sehr konkret, warum es passt].\n\n"
        "Eckdaten:\n"
        "  - Bestehende Plattform (web + mobile), aus Eigenmitteln gebaut\n"
        "  - 12-Monats-Pilot in Stuttgart + Partnerkommunen\n"
        "  - Förderbedarf: €85k (lean) bis €150k (voll)\n"
        "  - KPIs, Budget und Wirkungsmessung im beigefügten Deck\n"
        "  - Unabhängige Evaluation eingeplant\n\n"
        "Wir würden uns sehr über ein 30-minütiges Kennenlerngespräch "
        "freuen.\n\n"
        "Herzliche Grüße\n"
        "[Vorname Nachname], Mitgründer:in IndLokal\n"
        "hello@indlokal.de · indlokal.de",
        fill="F2F4F8",
    )

    add_h(doc, "Corporate CSR intro (EN, for Indian-IT majors)", 2)
    add_callout(
        doc,
        "Subject: IndLokal — supporting your Indian colleagues' integration into the Stuttgart region",
        "Dear [Name],\n\n"
        "I'm one of the founders of IndLokal, a digital platform that "
        "helps new Indian arrivals in Germany find local community, "
        "events, and trusted resources in their first 90 days.\n\n"
        "Many of your Stuttgart-region colleagues and their families "
        "are exactly our target users. We're launching a 12-month "
        "Stuttgart pilot in partnership with Stadt Stuttgart, "
        "Sindelfingen and Waiblingen, and we'd like to explore whether "
        "[Company] would like to be involved as a founding sponsor.\n\n"
        "What we'd offer in return:\n"
        "  - Visible co-branding in Stuttgart / Sindelfingen events\n"
        "  - A direct, useful integration tool for your incoming "
        "Indian hires and their families\n"
        "  - Quarterly impact reporting you can use for your own CSR / "
        "DEI reporting\n\n"
        "Sponsorship range: €5\\u201325k. 30 minutes to walk you "
        "through the platform and the pilot?\n\n"
        "Best regards,\n"
        "[Name], Co-founder, IndLokal\n"
        "hello@indlokal.de · indlokal.de",
        fill="F2F4F8",
    )

    # --- 9. Tracking
    add_h(doc, "Tracking sheet \u2014 maintain weekly", 1)
    add_p(
        doc,
        "Set up a simple Google Sheet / Notion table. Update weekly. "
        "When the first funder asks 'who else are you talking to?' you "
        "want a credible answer.",
    )
    add_table(
        doc,
        ["Funder", "Type", "Amount asked", "Status", "Next action", "Owner", "Date"],
        [
            ["Stadt Stuttgart Abt. Integration", "City", "Letter of Support",
             "(e.g. Email sent 25.04)", "Follow-up 02.05", "Dhiraj", "—"],
            ["Stadt Sindelfingen", "City", "LoS + €5\\u201315k Integrationsfonds",
             "Email sent", "Follow-up", "Jaya", "—"],
            ["Stadt Waiblingen", "City", "LoS + €5\\u201315k Integrationsfonds",
             "Email sent", "Follow-up", "Dhiraj", "—"],
            ["Robert Bosch Stiftung", "Foundation", "€50\\u201385k",
             "Researching contact", "Find prog. officer", "Jaya", "—"],
            ["Stiftung Mercator", "Foundation", "€85\\u2013150k",
             "Not started", "—", "—", "—"],
            ["BW Stiftung", "Foundation", "€85\\u2013150k",
             "Not started", "Check current call", "—", "—"],
            ["AMIF national (BAMF)", "Public", "€100\\u2013250k",
             "Not started", "Email regional coordinator", "—", "—"],
            ["ESF+ BW (Min. f. Soziales)", "Public", "€50\\u2013150k",
             "Not started", "Check current call", "—", "—"],
        ],
    )

    # --- 10. Quarterly cadence
    add_h(doc, "Quarterly cadence (after the first 3 months)", 1)
    add_bullets(
        doc,
        [
            "Submit at least 1 funder application per month, on average.",
            "Run 1 user-research touch per quarter (5 short interviews).",
            "Publish quarterly impact update — even before any grant lands. Builds credibility.",
            "Refresh the deck KPIs with actual numbers each quarter. The deck must be alive, not frozen.",
            "Keep an audit-ready folder: receipts, contracts, timesheets, decisions.",
        ],
    )

    # --- 11. What 'success' looks like in 6 months
    add_h(doc, "What success looks like at the 6-month mark", 1)
    add_checklist(
        doc,
        [
            "Legal entity (gUG / e.V. / gGmbH) registered with confirmed Gemeinn\u00fctzigkeit.",
            "Platform live in production with 30+ communities, 100+ resources.",
            "At least 2 Letters of Support in hand (city or institutional).",
            "At least 1 foundation cheque committed (\u20ac10\u201385k).",
            "At least 1 corporate sponsorship committed (\u20ac5\u201325k).",
            "AMIF or ESF+ BW application submitted to the next available call.",
            "Quarterly impact report #1 published with real numbers.",
        ],
    )

    add_callout(
        doc,
        "Final word",
        "These docs are tools, not the work. The work is conversations: "
        "with city officials, programme officers, community leaders, "
        "corporate CSR contacts. The docs make those conversations "
        "10x faster and 10x more credible. Ship the conversations.",
        fill="E8F0FF",
    )

    # closing
    doc.add_paragraph()
    disc = doc.add_paragraph()
    r = disc.add_run(
        "Working action plan for the IndLokal founding team. Update "
        "weekly. Programme names, contacts and call schedules change "
        "frequently — verify on the EU Funding & Tenders Portal, BAMF, "
        "L-Bank BW and each funder's site before applying."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
